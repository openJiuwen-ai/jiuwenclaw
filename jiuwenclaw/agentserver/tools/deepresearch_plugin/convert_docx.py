import base64
import logging
import re
import tempfile
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path

import pypandoc
import requests


logger = logging.getLogger(__name__)

TEXT_READ_ENCODINGS = ("utf-8-sig", "utf-8", "gb18030", "gbk")
DEFAULT_DOCX_FONT = "Microsoft YaHei"
NUMBERED_HEADING_RE = re.compile(
    r"^(?P<indent>\s{0,3})(?P<number>\d+(?:\.\d+)*)(?:\.\s+|\s+)(?P<title>.+?)\s*$"
)
SENTENCE_END_RE = re.compile(r"[。！？?!；;：:]$")
MERMAID_BLOCK_RE = re.compile(r"```mermaid\s*\n(.*?)```", re.DOTALL | re.IGNORECASE)


@dataclass(slots=True)
class MermaidRenderStats:
    total: int = 0
    success: int = 0
    failed: int = 0

try:
    import yaml
    YAML_AVAILABLE = True
except Exception:
    YAML_AVAILABLE = False

try:
    from PIL import Image, ImageEnhance
    PIL_AVAILABLE = True
except Exception:
    PIL_AVAILABLE = False

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_STYLE_AVAILABLE = True
except Exception:
    DOCX_STYLE_AVAILABLE = False


# =========================
# Pandoc
# =========================
def ensure_pandoc():
    try:
        pypandoc.get_pandoc_version()
    except OSError:
        logger.info("自动下载 Pandoc...")
        pypandoc.download_pandoc()


def read_text_with_fallback(path: Path) -> str:
    last_error: UnicodeDecodeError | None = None

    for encoding in TEXT_READ_ENCODINGS:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError as exc:
            last_error = exc

    raise UnicodeDecodeError(
        getattr(last_error, "encoding", "unknown"),
        getattr(last_error, "object", b""),
        getattr(last_error, "start", 0),
        getattr(last_error, "end", 0),
        f"无法正确解码文件：{path}",
    )


def _set_rfonts(rpr, font_name: str) -> None:
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)

    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), font_name)


def _set_run_font(run, font_name: str) -> None:
    run.font.name = font_name
    _set_rfonts(run.element.get_or_add_rPr(), font_name)


def _set_style_font(style, font_name: str) -> None:
    style.font.name = font_name
    _set_rfonts(style.element.get_or_add_rPr(), font_name)


def _apply_font_to_table(table, font_name: str) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run, font_name)
            for nested_table in cell.tables:
                _apply_font_to_table(nested_table, font_name)


def normalize_docx_fonts(docx_path: Path, *, font_name: str = DEFAULT_DOCX_FONT) -> None:
    if not DOCX_STYLE_AVAILABLE:
        logger.warning("python-docx 不可用，跳过 DOCX 字体统一。")
        return

    document = Document(docx_path)

    for style_name in (
        "Normal",
        "Title",
        "Subtitle",
        "Body Text",
        "Hyperlink",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Heading 4",
        "Heading 5",
        "Heading 6",
        "Heading 7",
        "Heading 8",
        "Heading 9",
    ):
        try:
            _set_style_font(document.styles[style_name], font_name)
        except KeyError:
            continue

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            _set_run_font(run, font_name)

    for table in document.tables:
        _apply_font_to_table(table, font_name)

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            for run in paragraph.runs:
                _set_run_font(run, font_name)
        for paragraph in section.footer.paragraphs:
            for run in paragraph.runs:
                _set_run_font(run, font_name)
        for table in section.header.tables:
            _apply_font_to_table(table, font_name)
        for table in section.footer.tables:
            _apply_font_to_table(table, font_name)

    document.save(docx_path)


# =========================
# Mermaid 配置与清洗
# =========================
DEFAULT_CONFIG = {
    "theme": "base",
    "look": "classic",
    "themeVariables": {
        "background": "#ffffff",
        "primaryTextColor": "#111827",
        "secondaryTextColor": "#111827",
        "tertiaryTextColor": "#111827",
        "lineColor": "#374151",
        "textColor": "#111827",
        "mainBkg": "#ffffff",
        "secondBkg": "#f9fafb",
        "tertiaryColor": "#ffffff",
        "xyChart": {
            "plotColorPalette": "#4338ca, #b91c1c, #047857, #b45309, #6d28d9"
        },
    },
}


def _deep_merge(base: dict, override: dict) -> dict:
    result = deepcopy(base)
    for k, v in (override or {}).items():
        if isinstance(v, dict) and isinstance(result.get(k), dict):
            result[k] = _deep_merge(result[k], v)
        else:
            result[k] = deepcopy(v)
    return result


def _extract_frontmatter(code: str):
    m = re.match(r"^\s*---\s*\n(.*?)\n---\s*\n?", code.strip(), flags=re.DOTALL)
    if m:
        return m.group(1), code.strip()[m.end():].strip()
    return "", code.strip()


def _dump_frontmatter(config_dict: dict) -> str:
    if YAML_AVAILABLE:
        text = yaml.safe_dump(
            {"config": config_dict},
            allow_unicode=True,
            sort_keys=False,
            default_flow_style=False,
        ).strip()
        return f"---\n{text}\n---\n"
    # 最小兜底：手写一个稳定配置块
    tv = config_dict.get("themeVariables", {})
    xy = tv.get("xyChart", {})
    return (
        "---\n"
        "config:\n"
        f"  theme: {config_dict.get('theme', 'base')}\n"
        f"  look: {config_dict.get('look', 'classic')}\n"
        "  themeVariables:\n"
        f"    background: '{tv.get('background', '#ffffff')}'\n"
        f"    primaryTextColor: '{tv.get('primaryTextColor', '#111827')}'\n"
        f"    secondaryTextColor: '{tv.get('secondaryTextColor', '#111827')}'\n"
        f"    tertiaryTextColor: '{tv.get('tertiaryTextColor', '#111827')}'\n"
        f"    lineColor: '{tv.get('lineColor', '#374151')}'\n"
        f"    textColor: '{tv.get('textColor', '#111827')}'\n"
        f"    mainBkg: '{tv.get('mainBkg', '#ffffff')}'\n"
        f"    secondBkg: '{tv.get('secondBkg', '#f9fafb')}'\n"
        f"    tertiaryColor: '{tv.get('tertiaryColor', '#ffffff')}'\n"
        "    xyChart:\n"
        f"      plotColorPalette: '{xy.get('plotColorPalette', '#4338ca, #b91c1c, #047857, #b45309, #6d28d9')}'\n"
        "---\n"
    )


def _build_merged_frontmatter(frontmatter: str, body: str) -> str:
    if not frontmatter:
        return _dump_frontmatter(DEFAULT_CONFIG) + body.strip()

    if not YAML_AVAILABLE:
        logger.warning("未安装 PyYAML，已有 frontmatter 时将保留原配置，不注入默认 Mermaid 配置。")
        return f"---\n{frontmatter.strip()}\n---\n{body.strip()}"

    try:
        parsed = yaml.safe_load(frontmatter) or {}
        if not isinstance(parsed, dict):
            parsed = {}
    except Exception as e:
        logger.warning("frontmatter 解析失败，将改用默认配置: %s", e)
        return _dump_frontmatter(DEFAULT_CONFIG) + body.strip()

    if "config" in parsed and isinstance(parsed["config"], dict):
        existing_config = parsed["config"]
    else:
        # 少数情况下用户直接把 config 项写在根层，尽量兼容
        existing_config = parsed if isinstance(parsed, dict) else {}

    merged_config = _deep_merge(DEFAULT_CONFIG, existing_config)

    # 一些高对比默认值，如果用户没设则保留默认；用户设了就尊重用户
    merged_config.setdefault("theme", "base")
    merged_config.setdefault("look", "classic")
    merged_config.setdefault("themeVariables", {})
    merged_config["themeVariables"].setdefault(
        "xyChart", {"plotColorPalette": "#4338ca, #b91c1c, #047857, #b45309, #6d28d9"}
    )
    merged_config["themeVariables"]["xyChart"].setdefault(
        "plotColorPalette", "#4338ca, #b91c1c, #047857, #b45309, #6d28d9"
    )

    return _dump_frontmatter(merged_config) + body.strip()


def clean_mermaid_code(code: str) -> str:
    code = code.strip()
    frontmatter, body = _extract_frontmatter(code)
    return _build_merged_frontmatter(frontmatter, body).strip()


# =========================
# 图片增强
# =========================
def enhance_image(image_path: str):
    if not PIL_AVAILABLE:
        return

    try:
        with Image.open(image_path) as original:
            if original.mode in ("RGBA", "LA"):
                bg = Image.new("RGBA", original.size, (255, 255, 255, 255))
                bg.alpha_composite(original.convert("RGBA"))
                img = bg.convert("RGB")
            else:
                img = original.convert("RGB")
        img = img.resize((img.width * 2, img.height * 2), Image.Resampling.LANCZOS)
        img = ImageEnhance.Contrast(img).enhance(1.24)
        img = ImageEnhance.Sharpness(img).enhance(1.35)
        img = ImageEnhance.Color(img).enhance(1.15)
        img.save(image_path, format="PNG", optimize=True)
    except Exception as e:
        logger.warning("图片增强失败: %s", e)


# =========================
# Mermaid 渲染
# =========================
def _valid_image_response(resp: requests.Response) -> bool:
    content_type = (resp.headers.get("Content-Type") or "").lower()
    if resp.status_code != 200:
        return False
    if not resp.content:
        return False
    if content_type and "image/" not in content_type:
        return False
    return True


def _save_failed_source(code: str, debug_base_path: Path, extra_text: str = ""):
    debug_base_path.parent.mkdir(parents=True, exist_ok=True)
    failed_src = debug_base_path.with_suffix(".mmd")
    failed_src.write_text(code, encoding="utf-8")
    logger.warning("已保存 Mermaid 源码: %s", failed_src)
    if extra_text:
        failed_log = debug_base_path.with_suffix(".error.txt")
        failed_log.write_text(extra_text, encoding="utf-8")
        logger.warning("已保存错误信息: %s", failed_log)


def _error_excerpt(resp: requests.Response) -> str:
    content_type = (resp.headers.get("Content-Type") or "").lower()
    excerpt = f"status={resp.status_code}\ncontent-type={content_type}\n"
    try:
        if "text/" in content_type or "json" in content_type or "xml" in content_type:
            excerpt += resp.text[:2000]
        else:
            excerpt += f"binary response, {len(resp.content)} bytes"
    except Exception:
        excerpt += "unable to decode response body"
    return excerpt


def render_mermaid(code: str, output_path: str, *, debug_base_path: Path | None = None) -> bool:
    code = clean_mermaid_code(code)
    last_error = ""
    output_file = Path(output_path)

    try:
        encoded = base64.urlsafe_b64encode(code.encode("utf-8")).decode("utf-8")
        url = f"https://mermaid.ink/img/{encoded}?bgColor=!white&type=png"
        r = requests.get(
            url,
            timeout=20,
            headers={"Accept": "image/png,image/*;q=0.9,*/*;q=0.8"},
        )
        if _valid_image_response(r):
            with output_file.open("wb") as f:
                f.write(r.content)
            enhance_image(str(output_file))
            return True
        last_error = "mermaid.ink\n" + _error_excerpt(r)
        logger.warning(
            "mermaid.ink 失败: status=%s, content-type=%s",
            r.status_code,
            r.headers.get("Content-Type"),
        )
    except Exception as e:
        last_error = f"mermaid.ink exception\n{e}"
        logger.exception("mermaid.ink 异常")

    try:
        url = "https://kroki.io/mermaid/png"
        r = requests.post(
            url,
            data=code.encode("utf-8"),
            headers={
                "Content-Type": "text/plain; charset=utf-8",
                "Accept": "image/png,image/*;q=0.9,*/*;q=0.8",
            },
            timeout=25,
        )
        if _valid_image_response(r):
            with output_file.open("wb") as f:
                f.write(r.content)
            enhance_image(str(output_file))
            return True
        last_error += "\n\nkroki\n" + _error_excerpt(r)
        logger.warning(
            "kroki 失败: status=%s, content-type=%s",
            r.status_code,
            r.headers.get("Content-Type"),
        )
    except Exception as e:
        last_error += f"\n\nkroki exception\n{e}"
        logger.exception("kroki 异常")

    _save_failed_source(code, debug_base_path or output_file, last_error)
    logger.warning(
        "在线 API 全部失败，可手动尝试以下网站:\n"
        "- https://www.mermaidonline.live/zh/mermaid-to-image\n"
        "- https://mermaid2img.com/zh-CN\n"
        "- https://mermaid.live/"
    )
    return False


# =========================
# 标题修复（保守版）
# =========================
def _neighbor_numbered_line(lines: list[str], index: int, *, reverse: bool) -> str | None:
    step = -1 if reverse else 1
    cursor = index + step

    while 0 <= cursor < len(lines):
        stripped = lines[cursor].strip()
        if stripped:
            return stripped
        cursor += step

    return None


def _should_promote_numbered_heading(lines: list[str], index: int, match: re.Match[str]) -> bool:
    title = match.group("title").strip()

    if len(title) > 80 or SENTENCE_END_RE.search(title):
        return False

    if index > 0 and lines[index - 1].strip():
        return False

    if index + 1 < len(lines) and lines[index + 1].strip():
        return False

    prev_nonempty = _neighbor_numbered_line(lines, index, reverse=True)
    next_nonempty = _neighbor_numbered_line(lines, index, reverse=False)
    if prev_nonempty and NUMBERED_HEADING_RE.match(prev_nonempty):
        return False
    if next_nonempty and NUMBERED_HEADING_RE.match(next_nonempty):
        return False

    return True


def normalize_headings(content: str) -> str:
    content = content.replace("\r\n", "\n").replace("\r", "\n")

    lines = content.split("\n")
    out = []
    in_code_block = False

    for index, raw_line in enumerate(lines):
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code_block = not in_code_block
            out.append(line)
            continue

        if in_code_block:
            out.append(line)
            continue

        if not stripped:
            out.append("")
            continue

        m_hash = re.match(r"^(#{1,6})\s*(.+?)\s*$", stripped)
        if m_hash:
            hashes = m_hash.group(1)
            title = m_hash.group(2).strip()

            if out and out[-1] != "":
                out.append("")
            out.append(f"{hashes} {title}")
            out.append("")
            continue

        m_num = NUMBERED_HEADING_RE.match(line)
        if m_num and _should_promote_numbered_heading(lines, index, m_num):
            numbering = m_num.group("number")
            title = m_num.group("title").strip()
            level = min(numbering.count(".") + 1, 6)
            heading = "#" * level + " " + f"{numbering} {title}"

            if out and out[-1] != "":
                out.append("")
            out.append(heading)
            out.append("")
            continue

        out.append(line)

    result = "\n".join(out)
    result = re.sub(r"\n{3,}", "\n\n", result)
    return result.strip() + "\n"


# =========================
# Mermaid 替换
# =========================
def replace_mermaid_blocks(content: str, tmp_dir: Path, *, debug_dir: Path, debug_stem: str):
    stats = MermaidRenderStats()

    def repl(match):
        stats.total += 1
        block_index = stats.total - 1
        code = match.group(1).strip()
        img_name = f"mermaid_{block_index}.png"
        img_path = tmp_dir / img_name
        debug_base_path = debug_dir / f"{debug_stem}_mermaid_{block_index}"

        if render_mermaid(code, str(img_path), debug_base_path=debug_base_path):
            stats.success += 1
            return f"\n\n![diagram]({img_name})\n\n"

        logger.warning("Mermaid 渲染失败，保留原代码块")
        stats.failed += 1
        return match.group(0)

    new_content = MERMAID_BLOCK_RE.sub(repl, content)
    return new_content, stats


# =========================
# 主流程
# =========================
def convert_md_to_docx(md_path: str, docx_path: str):
    ensure_pandoc()

    md_path = Path(md_path).resolve()
    docx_path = Path(docx_path).resolve()

    if not md_path.exists():
        raise FileNotFoundError(f"找不到 Markdown 文件: {md_path}")

    docx_path.parent.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        content = read_text_with_fallback(md_path)
        content = normalize_headings(content)
        content, mermaid_stats = replace_mermaid_blocks(
            content,
            tmp_dir,
            debug_dir=docx_path.parent,
            debug_stem=docx_path.stem,
        )

        temp_md = tmp_dir / "temp.md"
        temp_md.write_text(content, encoding="utf-8")

        pypandoc.convert_file(
            str(temp_md),
            "docx",
            outputfile=str(docx_path),
            extra_args=[
                "--from=gfm",
                "--resource-path",
                str(tmp_dir),
            ],
        )
        normalize_docx_fonts(docx_path)

        logger.info("转换完成: %s", docx_path)
        logger.info("DOCX 字体已统一为: %s", DEFAULT_DOCX_FONT)
        logger.info(
            "Mermaid 图表统计: 总数=%s, 成功=%s, 失败=%s",
            mermaid_stats.total,
            mermaid_stats.success,
            mermaid_stats.failed,
        )


if __name__ == "__main__":
    try:
        convert_md_to_docx("input.md", "output.docx")
    except Exception:
        logger.exception("转换失败")
        raise
