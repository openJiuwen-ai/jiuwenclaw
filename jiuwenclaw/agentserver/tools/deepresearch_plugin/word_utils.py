# Copyright (c) Huawei Technologies Co., Ltd. 2025. All rights reserved.
from __future__ import annotations

import base64
import io
import re
from dataclasses import dataclass, replace
from pathlib import Path
from urllib.parse import unquote, urlparse

from bs4 import BeautifulSoup, NavigableString
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Pt
from latex2mathml.converter import convert as latex_to_mathml
from mathml2omml import convert as mathml_to_omml

MAX_HTML_BLOCK_DEPTH = 100
HEADING_TAGS = frozenset(f"h{i}" for i in range(1, 10))
REMOTE_IMAGE_SCHEMES = frozenset({"http", "https"})
LATEX_TOKEN_RE = re.compile(r"(\$\$.*?\$\$|\\\(.*?\\\))", re.DOTALL)
LATEX_GROUPED_COMMANDS_WITH_POWER = frozenset({"binom", "frac"})
LATEX_NORMALIZATION_MAX_PASSES = 8
LATEX_ALIGNMENT_ENV_RE = re.compile(
    r"\\begin\{(?P<env>align\*?|aligned|split|gathered)\}"
    r"(?P<body>.*?)"
    r"\\end\{(?P=env)\}",
    re.DOTALL,
)
HTML_FORMATTING_WHITESPACE_RE = re.compile(r"[ \t]*\n[ \t]*")
DOCX_LIST_LEVELS = 9
DOCX_BULLET_SYMBOLS = ("•", "○", "▪")
HEADING_STYLE_SIZES = {
    "Heading 1": 22,
    "Heading 2": 18,
    "Heading 3": 16,
    "Heading 4": 14,
    "Heading 5": 12,
    "Heading 6": 11,
    "Heading 7": 11,
    "Heading 8": 11,
    "Heading 9": 11,
}


@dataclass(frozen=True)
class HtmlToDocContext:
    style_dict: dict[str, str]
    base_path: Path | None = None
    max_image_width: int | None = None
    max_depth: int = MAX_HTML_BLOCK_DEPTH
    style_r_fonts: object | None = None
    bold: bool = False
    italic: bool = False
    underline: bool = False
    in_code: bool = False
    superscript: bool = False


@dataclass(frozen=True)
class HtmlBlockState:
    depth: int = 0
    list_depth: int = 0
    list_num_id: int | None = None
    list_tag: str | None = None


def _docx_paragraph_p(paragraph):
    return paragraph._p  # pylint: disable=protected-access


def _append_word_list_level(abstract_num, level: int, *, ordered: bool) -> None:
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), str(level))

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    lvl.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(
        qn("w:val"),
        f"%{level + 1}." if ordered else DOCX_BULLET_SYMBOLS[level % len(DOCX_BULLET_SYMBOLS)],
    )
    lvl.append(lvl_text)

    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), str(360 * (level + 1)))
    indent.set(qn("w:hanging"), "180")
    p_pr.append(indent)
    lvl.append(p_pr)
    abstract_num.append(lvl)


def _create_word_list_numbering(doc, *, ordered: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    abstract_num_id = max(abstract_ids, default=-1) + 1

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_num_id))
    multi_level_type = OxmlElement("w:multiLevelType")
    multi_level_type.set(qn("w:val"), "multilevel")
    abstract_num.append(multi_level_type)
    for level in range(DOCX_LIST_LEVELS):
        _append_word_list_level(abstract_num, level, ordered=ordered)
    numbering.insert(len(abstract_ids), abstract_num)

    num = numbering.add_num(abstract_num_id)
    return int(num.get(qn("w:numId")))


def _apply_word_list_numbering(paragraph, num_id: int, level: int) -> None:
    num_pr = _docx_paragraph_p(paragraph).get_or_add_pPr().get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(min(level, DOCX_LIST_LEVELS - 1)))
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_element)


def _get_style_by_tag(tag_name: str, style_dict: dict[str, str], doc):
    key = tag_name
    if tag_name in HEADING_TAGS:
        key = f"heading{tag_name[1:]}"
    elif tag_name == "table":
        key = "table"
    elif tag_name not in style_dict:
        key = "paragraph"

    style_name = style_dict.get(key) or style_dict.get("default") or "Normal"
    try:
        return doc.styles[style_name]
    except KeyError:
        return doc.styles["Normal"]


def _apply_style_font_on_run(run, style_r_fonts) -> None:
    if style_r_fonts is None:
        return
    r_pr = run.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        val = style_r_fonts.get(qn(attr))
        if val:
            r_fonts.set(qn(attr), val)


def _add_text_run(paragraph, text: str, context: HtmlToDocContext):
    run = paragraph.add_run(text)
    _apply_style_font_on_run(run, context.style_r_fonts)
    if context.bold:
        run.bold = True
    if context.italic:
        run.italic = True
    if context.underline:
        run.underline = True
    if context.in_code:
        run.font.name = "Consolas"
    if context.superscript:
        run.font.superscript = True
    return run


def _latex_token_content(token: str) -> str:
    if token.startswith("$$") and token.endswith("$$"):
        return token[2:-2].strip()
    if token.startswith(r"\(") and token.endswith(r"\)"):
        return token[2:-2].strip()
    return token.strip()


def _add_latex_math(paragraph, latex_token: str, context: HtmlToDocContext) -> None:
    latex = _latex_token_content(latex_token)
    if not latex:
        return
    try:
        _append_child_to_paragraph(paragraph, _latex_to_omml(latex, context))
    except ValueError:
        _add_text_run(paragraph, latex_token, context)


def _process_text_with_latex(paragraph, text: str, context: HtmlToDocContext) -> None:
    if context.in_code:
        _add_text_run(paragraph, text, context)
        return

    cursor = 0
    for match in LATEX_TOKEN_RE.finditer(text):
        if match.start() > cursor:
            _add_text_run(paragraph, text[cursor:match.start()], context)
        _add_latex_math(paragraph, match.group(0), context)
        cursor = match.end()
    if cursor < len(text):
        _add_text_run(paragraph, text[cursor:], context)


def _latex_to_omml(latex: str, context: HtmlToDocContext):
    """Create a real Word math element for LaTeX content."""
    latex = _normalize_latex_for_omml(latex)
    if not latex:
        raise ValueError("empty latex")

    try:
        mathml = latex_to_mathml(latex)
        omml = mathml_to_omml(mathml)
        return parse_xml(omml.replace("<m:oMath", f"<m:oMath {nsdecls('m')}", 1))
    except Exception as exc:
        raise ValueError(f"Unable to convert LaTeX to OMML: {latex}") from exc


def _normalize_latex_for_omml(latex: str) -> str:
    """Normalize valid LaTeX forms that Word math import handles more reliably."""
    previous = _strip_latex_alignment_markers(latex)
    for _ in range(LATEX_NORMALIZATION_MAX_PASSES):
        current = _wrap_grouped_command_powers(previous)
        if current == previous:
            return current
        previous = current
    return previous


def _strip_latex_alignment_markers(latex: str) -> str:
    """Remove unescaped alignment markers from LaTeX alignment environments."""

    def _strip_environment(match: re.Match[str]) -> str:
        env = match.group("env")
        body = _strip_unescaped_latex_char(match.group("body"), "&")
        return rf"\begin{{{env}}}{body}\end{{{env}}}"

    return LATEX_ALIGNMENT_ENV_RE.sub(_strip_environment, latex)


def _strip_unescaped_latex_char(text: str, target: str) -> str:
    parts: list[str] = []
    index = 0
    while index < len(text):
        char = text[index]
        if char == "\\" and index + 1 < len(text):
            parts.append(text[index:index + 2])
            index += 2
            continue
        if char != target:
            parts.append(char)
        index += 1
    return "".join(parts)


def _wrap_grouped_command_powers(latex: str) -> str:
    parts: list[str] = []
    cursor = 0
    index = 0

    while index < len(latex):
        match = re.search(r"\\([A-Za-z]+)", latex[index:])
        if match is None:
            break

        command_start = index + match.start()
        command_end = index + match.end()
        command_name = match.group(1)
        if command_name not in LATEX_GROUPED_COMMANDS_WITH_POWER:
            index = command_end
            continue

        first_group_end = _find_latex_group_end(latex, command_end)
        if first_group_end is None:
            index = command_end
            continue
        second_group_end = _find_latex_group_end(latex, first_group_end + 1)
        if second_group_end is None:
            index = first_group_end + 1
            continue

        power_end = _find_latex_power_end(latex, second_group_end + 1)
        if power_end is None:
            index = command_end
            continue

        parts.append(latex[cursor:command_start])
        parts.append("{")
        parts.append(latex[command_start:second_group_end + 1])
        parts.append("}")
        parts.append(latex[second_group_end + 1:power_end])
        cursor = power_end
        index = power_end

    if not parts:
        return latex

    parts.append(latex[cursor:])
    return "".join(parts)


def _find_latex_group_end(text: str, open_index: int) -> int | None:
    if open_index >= len(text) or text[open_index] != "{":
        return None

    depth = 0
    index = open_index
    while index < len(text):
        char = text[index]
        if char == "\\":
            index += 2
            continue
        if char == "{":
            depth += 1
        elif char == "}":
            depth -= 1
            if depth == 0:
                return index
        index += 1
    return None


def _find_latex_power_end(text: str, caret_index: int) -> int | None:
    if caret_index >= len(text) or text[caret_index] != "^":
        return None

    value_start = caret_index + 1
    if value_start >= len(text):
        return None
    if text[value_start] == "{":
        group_end = _find_latex_group_end(text, value_start)
        return None if group_end is None else group_end + 1
    if text[value_start].isalnum():
        return value_start + 1
    return None


def _apply_r_fonts_to_r_pr(r_pr, style_r_fonts) -> None:
    if style_r_fonts is None:
        return
    r_fonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        val = style_r_fonts.get(qn(attr))
        if val:
            r_fonts.set(qn(attr), val)
    r_pr.append(r_fonts)


def _append_child_to_paragraph(paragraph, child_element) -> None:
    getattr(paragraph, "_p").append(child_element)


def _add_hyperlink(paragraph, url: str, text: str, context: HtmlToDocContext) -> None:
    r_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    _apply_r_fonts_to_r_pr(r_pr, context.style_r_fonts)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    if context.superscript:
        vert_align = OxmlElement("w:vertAlign")
        vert_align.set(qn("w:val"), "superscript")
        r_pr.append(vert_align)

    run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    _append_child_to_paragraph(paragraph, hyperlink)


def _is_relative_to(path: Path, base_path: Path) -> bool:
    try:
        path.relative_to(base_path)
        return True
    except ValueError:
        return False


def _resolve_local_image(src: str, base_path: Path | None) -> Path | None:
    if base_path is None or not src:
        return None

    parsed = urlparse(src)
    if parsed.scheme in REMOTE_IMAGE_SCHEMES or parsed.scheme == "data":
        return None

    raw_path = unquote(parsed.path or src)
    image_path = Path(raw_path)
    base_resolved = base_path.resolve()
    if image_path.is_absolute():
        candidate = image_path.resolve()
    else:
        candidate = (base_resolved / image_path).resolve()

    if not _is_relative_to(candidate, base_resolved):
        return None
    return candidate if candidate.exists() else None


def _fit_inline_shape_to_width(inline_shape, max_width) -> None:
    if max_width is None or inline_shape.width <= max_width:
        return
    if inline_shape.width == 0:
        inline_shape.width = max_width
        return
    scale = max_width / inline_shape.width
    inline_shape.width = max_width
    inline_shape.height = int(inline_shape.height * scale)


def _process_image(paragraph, src: str, context: HtmlToDocContext) -> None:
    if src.startswith("data:image/") and ";base64," in src:
        _, b64data = src.split(",", 1)
        image_bytes = base64.b64decode(b64data)
        run = paragraph.add_run()
        inline_shape = run.add_picture(io.BytesIO(image_bytes))
        _fit_inline_shape_to_width(inline_shape, context.max_image_width)
        return

    image_path = _resolve_local_image(src, context.base_path)
    if image_path is None:
        return

    run = paragraph.add_run()
    inline_shape = run.add_picture(str(image_path))
    _fit_inline_shape_to_width(inline_shape, context.max_image_width)


def _process_inline(paragraph, node, context: HtmlToDocContext) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if not text:
            return
        if "\n" in text:
            if not text.strip():
                return
            text = HTML_FORMATTING_WHITESPACE_RE.sub(" ", text)
        _process_text_with_latex(paragraph, text, context)
        return

    if node.name == "br":
        paragraph.add_run().add_break()
        return

    if node.name == "img":
        _process_image(paragraph, node.get("src") or "", context)
        return

    if node.name == "a":
        href = node.get("href")
        text = node.get_text(strip=False)
        if href and text:
            _add_hyperlink(paragraph, href, text, context)
        return

    if node.name == "sup":
        for child in node.contents:
            _process_inline(paragraph, child, replace(context, superscript=True))
        return

    if node.name in ("strong", "b", "em", "i", "u", "code"):
        inline_context = context
        if node.name in ("strong", "b"):
            inline_context = replace(inline_context, bold=True)
        elif node.name in ("em", "i"):
            inline_context = replace(inline_context, italic=True)
        elif node.name == "u":
            inline_context = replace(inline_context, underline=True)
        elif node.name == "code":
            inline_context = replace(inline_context, in_code=True)
        for child in node.contents:
            _process_inline(paragraph, child, inline_context)
        return

    for child in node.contents:
        _process_inline(paragraph, child, context)


def _add_para_and_apply_style(doc, element, context: HtmlToDocContext) -> None:
    style = _get_style_by_tag(element.name, context.style_dict, doc)
    if element.name == "p" and any(getattr(child, "name", None) == "img" for child in element.contents):
        _add_paragraph_with_split_images(doc, element, context, style)
        return

    paragraph = doc.add_paragraph(style=style)
    style_r_fonts = style.element.get_or_add_rPr().find(qn("w:rFonts"))
    for child in element.contents:
        _process_inline(paragraph, child, replace(context, style_r_fonts=style_r_fonts))


def _add_paragraph_with_split_images(doc, element, context: HtmlToDocContext, style) -> None:
    style_r_fonts = style.element.get_or_add_rPr().find(qn("w:rFonts"))
    paragraph = None

    def _paragraph():
        nonlocal paragraph
        if paragraph is None:
            paragraph = doc.add_paragraph(style=style)
        return paragraph

    for child in element.contents:
        if getattr(child, "name", None) == "img":
            image_paragraph = doc.add_paragraph(style=style)
            _process_inline(image_paragraph, child, replace(context, style_r_fonts=style_r_fonts))
            paragraph = None
            continue

        if isinstance(child, NavigableString) and not str(child).strip() and paragraph is None:
            continue
        if getattr(child, "name", None) == "br" and paragraph is None:
            continue

        _process_inline(_paragraph(), child, replace(context, style_r_fonts=style_r_fonts))


def _add_text_paragraph(doc, text: str, context: HtmlToDocContext) -> None:
    if text.strip():
        paragraph_style = _get_style_by_tag("p", context.style_dict, doc)
        doc.add_paragraph(text.strip(), style=paragraph_style)


def _add_list_item(
    doc,
    li_element,
    context: HtmlToDocContext,
    state: HtmlBlockState,
    *,
    list_num_id: int,
    list_tag: str,
) -> None:
    style = _get_style_by_tag("p", context.style_dict, doc)
    paragraph = doc.add_paragraph(style=style)
    _apply_word_list_numbering(paragraph, list_num_id, state.list_depth)
    style_r_fonts = style.element.get_or_add_rPr().find(qn("w:rFonts"))
    child_context = replace(context, style_r_fonts=style_r_fonts)

    for child in li_element.contents:
        if getattr(child, "name", None) in ("ul", "ol"):
            _process_block_element(
                doc,
                child,
                context,
                replace(
                    state,
                    depth=state.depth + 1,
                    list_depth=state.list_depth + 1,
                    list_num_id=list_num_id,
                    list_tag=list_tag,
                ),
            )
        else:
            _process_inline(paragraph, child, child_context)


def _add_html_table_to_doc(doc, element, context: HtmlToDocContext) -> None:
    rows = element.find_all("tr")
    if not rows:
        return

    max_cols = max(len(row.find_all(["th", "td"], recursive=False)) for row in rows)
    if max_cols == 0:
        return

    table = doc.add_table(rows=len(rows), cols=max_cols)
    table_style_name = context.style_dict.get("table")
    if table_style_name and table_style_name in doc.styles:
        table.style = doc.styles[table_style_name]

    paragraph_style = _get_style_by_tag("p", context.style_dict, doc)
    style_r_fonts = paragraph_style.element.get_or_add_rPr().find(qn("w:rFonts"))
    cell_context = replace(context, style_r_fonts=style_r_fonts)

    for row_index, row in enumerate(rows):
        cells = row.find_all(["th", "td"], recursive=False)
        for col_index, cell in enumerate(cells[:max_cols]):
            target_cell = table.cell(row_index, col_index)
            paragraph = target_cell.paragraphs[0]
            paragraph.style = paragraph_style
            for child in cell.contents:
                _process_inline(paragraph, child, cell_context)


def _process_pre_block(doc, element, context: HtmlToDocContext) -> None:
    paragraph_style = _get_style_by_tag("p", context.style_dict, doc)
    paragraph = doc.add_paragraph(style=paragraph_style)
    run = paragraph.add_run(element.get_text())
    run.font.name = "Consolas"


def _process_block_element(
    doc,
    element,
    context: HtmlToDocContext,
    state: HtmlBlockState = HtmlBlockState(),
) -> None:
    if isinstance(element, NavigableString):
        _add_text_paragraph(doc, str(element), context)
        return
    if element.name is None:
        return
    if state.depth >= context.max_depth:
        _add_text_paragraph(doc, element.get_text(strip=True), context)
        return

    if element.name in HEADING_TAGS or element.name in ("p", "blockquote"):
        _add_para_and_apply_style(doc, element, context)
        return

    if element.name == "pre":
        _process_pre_block(doc, element, context)
        return

    if element.name == "table":
        _add_html_table_to_doc(doc, element, context)
        return

    if element.name in ("ul", "ol"):
        current_list_num_id = (
            state.list_num_id
            if state.list_num_id is not None and state.list_tag == element.name
            else _create_word_list_numbering(doc, ordered=element.name == "ol")
        )
        for li_element in element.find_all("li", recursive=False):
            _add_list_item(
                doc,
                li_element,
                context,
                state,
                list_num_id=current_list_num_id,
                list_tag=element.name,
            )
        return

    if element.name == "hr":
        doc.add_paragraph("")
        return

    if element.name in ("div", "section", "article", "body", "html"):
        for child in element.children:
            _process_block_element(
                doc,
                child,
                context,
                replace(
                    state,
                    depth=state.depth + 1,
                    list_num_id=None,
                    list_tag=None,
                ),
            )
        return

    _add_para_and_apply_style(doc, element, context)


def _get_available_page_width(doc) -> int:
    section = doc.sections[0]
    return section.page_width - section.left_margin - section.right_margin


def html_to_doc(doc, html: str, style_dict: dict[str, str], base_path: str | Path | None = None) -> None:
    """Convert HTML content into an existing python-docx document."""
    soup = BeautifulSoup(html, "html.parser")
    container = soup.find("div", class_="report-container")
    if container is None:
        container = soup.body or soup

    context = HtmlToDocContext(
        style_dict=style_dict,
        base_path=Path(base_path).resolve() if base_path is not None else None,
        max_image_width=_get_available_page_width(doc),
    )
    for element in container.children:
        _process_block_element(doc, element, context)


def set_global_styles(doc, font_name: str = "微软雅黑", font_size: int = 11, line_spacing: float = 1.15) -> None:
    """Set readable global fonts and compact paragraph spacing for a DOCX document."""
    for style in doc.styles:
        if getattr(style, "type", None) != 1:
            continue

        style.font.name = font_name
        style.font.size = Pt(HEADING_STYLE_SIZES.get(style.name, font_size))
        if style.name in HEADING_STYLE_SIZES:
            style.font.bold = True
        style.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font_name)

        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph_format.line_spacing = line_spacing
        paragraph_format.space_after = Pt(6)
