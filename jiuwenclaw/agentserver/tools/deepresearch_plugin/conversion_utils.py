# Copyright (c) Huawei Technologies Co., Ltd. 2025. All rights reserved.
from __future__ import annotations

import html
import logging
import re
from pathlib import Path

import unicodedata

logger = logging.getLogger(__name__)

TEXT_READ_ENCODINGS = ("utf-8-sig", "utf-8", "gb18030", "gbk")
DEFAULT_DOCX_FONT = "Microsoft YaHei"
DOCX_IMAGE_UPSCALE_FACTOR = 3
DOCX_IMAGE_CONTRAST = 1.28
DOCX_IMAGE_SHARPNESS = 1.45
DOCX_IMAGE_COLOR = 1.08
SAFE_FILENAME_RE = re.compile(r"[^\w.-]+", re.UNICODE)
NUMBERED_HEADING_RE = re.compile(
    r"^(?P<indent>\s{0,3})(?P<number>\d+(?:\.\d+)*)(?:\.\s+|\s+)(?P<title>.+?)\s*$"
)
LIST_ITEM_RE = re.compile(r"^\s{0,3}(?:[-*+]\s+|\d+\.\s+)")
BLOCK_MATH_RE = re.compile(r"\$\$.+?\$\$", re.DOTALL)
SIMPLE_VARIABLE_MATH_RE = re.compile(r"[A-Za-z\u0370-\u03ff]")
VARIABLE_ATOM_PATTERN = (
    r"(?:\\[A-Za-z]+|[A-Za-z\u0370-\u03ff](?:_\{?[\w\\]+\}?|\^\{?[\w\\]+\}?)?)"
)
COMPACT_VARIABLE_MATH_RE = re.compile(
    r"[A-Za-z\u0370-\u03ff]{2,}(?:_\{?[\w\\]+\}?|\^\{?[\w\\]+\}?)*"
)
ALNUM_VARIABLE_MATH_RE = re.compile(
    r"[A-Za-z\u0370-\u03ff]+\d+(?:_\{?[\w\\]+\}?|\^\{?[\w\\]+\}?)*"
)
VARIABLE_LIST_ELLIPSIS_PATTERN = r"(?:\.{3}|\u2026|\\ldots|\\cdots)"
VARIABLE_LIST_PATTERN = (
    rf"{VARIABLE_ATOM_PATTERN}(?:\s*,\s*{VARIABLE_ATOM_PATTERN})+"
    rf"(?:\s*,\s*{VARIABLE_LIST_ELLIPSIS_PATTERN})?"
)
VARIABLE_LIST_MATH_RE = re.compile(VARIABLE_LIST_PATTERN)
PAREN_VARIABLE_LIST_MATH_RE = re.compile(rf"\(\s*{VARIABLE_LIST_PATTERN}\s*\)")
MATH_NAME_PATTERN = (
    r"(?:\\?[A-Za-z\u0370-\u03ff]+|[0-9]+|\\[A-Za-z]+)"
    r"(?:_\{?[\w\\]+\}?|\^\{?[\w\\]+\}?|\\[A-Za-z]+)*"
)
MATH_FUNCTION_CALL_PATTERN = rf"{MATH_NAME_PATTERN}\s*[\(\[][^)\]\n]+[\)\]]"
MATH_SPACING_PATTERN = r"(?:\s+|\\[,;:! ]|\\quad|\\qquad)"
MATH_DELIMITED_CONTENT_PATTERN = r"[^|\n]+"
MATH_FEATURE_RE = re.compile(
    r"(\\[A-Za-z]+|[_^={}]|[+\-*/!%]|[\u00d7\u00f7\u2212\u2211\u220f\u221a\u2264\u2265\u2260\u2248<>])"
)
STRONG_NUMERIC_MATH_FEATURE_RE = re.compile(
    r"(\\[A-Za-z]+|[_^{}]|[+\-*/!%]|[\u00d7\u00f7\u2212=<>])"
)
MATH_FUNCTION_CALL_RE = re.compile(MATH_FUNCTION_CALL_PATTERN)
PRIME_FUNCTION_CALL_RE = re.compile(rf"{MATH_NAME_PATTERN}'{{1,3}}\s*[\(\[][^)\]\n]+[\)\]]")
FUNCTION_CALL_SEQUENCE_RE = re.compile(
    rf"{MATH_FUNCTION_CALL_PATTERN}(?:{MATH_SPACING_PATTERN}{MATH_FUNCTION_CALL_PATTERN})+"
)
ABSOLUTE_VALUE_MATH_RE = re.compile(rf"\|{MATH_DELIMITED_CONTENT_PATTERN}\|")
NORM_MATH_RE = re.compile(
    rf"(?:\|\|{MATH_DELIMITED_CONTENT_PATTERN}\|\||"
    rf"\\\|{MATH_DELIMITED_CONTENT_PATTERN}\\\||"
    rf"\\lVert\s*.+?\s*\\rVert)"
)
PERCENT_MATH_RE = re.compile(rf"(?:[+-]?\d+(?:\.\d+)?|{MATH_NAME_PATTERN})\s*\\?%")
EQUATION_REFERENCE_RE = re.compile(r"\((?:Eq\.?|Equation)\s*\d+[A-Za-z]?\)", re.IGNORECASE)
NUMERIC_TUPLE_RE = re.compile(r"\(\s*[+-]?\d+(?:\.\d+)?(?:\s*,\s*[+-]?\d+(?:\.\d+)?)+\s*\)")
PRIME_VARIABLE_RE = re.compile(
    r"(?:\\[A-Za-z]+|[A-Za-z\u0370-\u03ff])(?:_\{?[\w\\]+\}?|\^\{?[\w\\]+\}?)*'+"
)
PLAIN_CURRENCY_RE = re.compile(
    r"\d{1,3}(?:,\d{3})*(?:\.\d+)?"
    r"(?:\s*(?:[-\u2013\u2014]|to)\s*\d{1,3}(?:,\d{3})*(?:\.\d+)?)?"
)
MATH_PLACEHOLDER = "\x00MATH{}\x00"
MATH_PLACEHOLDER_RE = re.compile(r"\x00MATH(\d+)\x00")
INLINE_MATH_DOLLAR_RE = re.compile(r"^\$(?!\$)(.*?)(?<!\$)\$$", re.DOTALL)
HTML_TAG_RE = re.compile(r"</?[A-Za-z][A-Za-z0-9:-]*(?:\s+[^<>]*)?>")
MARKDOWN_LINK_RE = re.compile(r"!?\[[^\]\n]+\]\([^)]+\)")
CODE_PLACEHOLDER = "\x00CODE{}\x00"
CODE_PLACEHOLDER_RE = re.compile(r"\x00CODE(\d+)\x00")
FENCED_CODE_RE = re.compile(
    r"(?ms)^(?P<indent>[ \t]{0,3})(?P<fence>`{3,}|~{3,})[^\n]*\n.*?\n(?P=indent)(?P=fence)[ \t]*$"
)
INLINE_CODE_RE = re.compile(r"`+[^`\n]*`+")
LIST_ITEM_WITH_INDENT_RE = re.compile(
    r"^(?P<indent>[ \t]*)(?P<marker>(?:[-*+]\s+|\d+\.\s+).*)$"
)
INDENTED_LIST_ITEM_RE = re.compile(r"^(?P<indent>[ \t]{4,})(?P<marker>(?:[-*+]\s+|\d+\.\s+).*)$")
MARKDOWN_IMAGE_LINE_RE = re.compile(r"^[ \t]*!\[[^\]]*\]\([^)]+\)[ \t]*$")
LEGACY_FONT_CAPTION_PREFIX_RE = re.compile(
    r'^[ \t]*<font\b[^>]*\bsize\s*=\s*["\']?2["\']?[^>]*>',
    flags=re.IGNORECASE,
)
INDENTED_HTML_BLOCK_END_RE = re.compile(
    r"^[ \t]+</(?:div|section|article|main)>[ \t]*$",
    flags=re.IGNORECASE,
)
SENTENCE_END_RE = re.compile(r"[。！？?!…]$")

CITATION_RE = re.compile(r"\[\[(\d+)\]\]\((https?://[^\s)]+(?:\([^\s)]+\)[^\s)]*)*)\)")
CHECKED_CITATION_RE = re.compile(
    r"\[\s*checked_citation:\s*\d+\s*\](\[\[\d+\]\]\((?:[^()]|\([^()]*\))*\))"
)
LEGACY_CITATION_RE = re.compile(r"\[\s*citation:\s*\d+\s*\]")
REFERENCE_LINE_RE = re.compile(r"^(?P<indent>\s*)\[(\d+)\]\.\s+(.*)$", re.MULTILINE)
CENTER_CAPTION_RE = re.compile(
    r'<div\s+style="text-align:\s*center;?">',
    flags=re.IGNORECASE,
)
LEGACY_FONT_CAPTION_LINE_RE = re.compile(
    r'^(?P<indent>[ \t]*)<font\b[^>]*\bsize\s*=\s*["\']?2["\']?[^>]*>(?P<body>.*?)</font>'
    r'(?P<citations>(?:<sup class="citation">.*?</sup>)*)[ \t]*$',
    flags=re.IGNORECASE | re.MULTILINE,
)
EXTERNAL_LINK_RE = re.compile(
    r'<a\s+([^>]*?)href="(https?://[^"]+)"(?![^>]*\btarget=)([^>]*)>',
    flags=re.IGNORECASE,
)
CITATION_ANCHOR_RE = re.compile(
    r'(?<!<sup class="citation">)(<a\b[^>]*href="https?://[^"]+"[^>]*>\[(\d+)\]</a>)',
    flags=re.IGNORECASE,
)
MERMAID_BLOCK_WITH_ADJACENT_CAPTION_RE = re.compile(
    r'(?is)```[ \t]*mermaid[ \t]*\r?\n(.*?)\r?\n```[ \t]*'
    r'(?:'
    r'(?:\r?\n[ \t]*)+'
    r'<div\b(?=[^>]*\b(?:style="text-align:\s*center;?"|class="figure-caption"))[^>]*>'
    r'.*?'
    r'</div>[ \t]*'
    r')?'
)

try:
    from PIL import Image, ImageEnhance

    PIL_AVAILABLE = True
except Exception:
    PIL_AVAILABLE = False

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    DOCX_STYLE_AVAILABLE = True
except Exception:
    DOCX_STYLE_AVAILABLE = False


def read_text_with_fallback(path: Path) -> str:
    """Read text file content with multiple encoding fallback attempts.
    
    Attempts to decode the file using a prioritized list of encodings
    (utf-8-sig, utf-8, gb18030, gbk) to handle various text file formats,
    particularly those originating from different locales.
    
    Args:
        path: Path object pointing to the text file to read.
    
    Returns:
        The decoded text content of the file as a string.
    
    Raises:
        UnicodeDecodeError: If none of the fallback encodings can decode the file.
            The error includes the path that failed to decode.
    
    Note:
        Encoding order prioritizes UTF-8 variants first, then Chinese encodings
        (gb18030, gbk) for compatibility with documents from Chinese sources.
    """
    last_error: UnicodeDecodeError | None = None

    for encoding in TEXT_READ_ENCODINGS:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError as exc:
            last_error = exc

    raise UnicodeDecodeError(
        getattr(last_error, "encoding", "unknown"),
        getattr(last_error, "object", b""),
        getattr(last_error, "start", 0),
        getattr(last_error, "end", 0),
        f"Unable to decode text file: {path}",
    )


def make_safe_filename_component(value: str, *, default: str = "document") -> str:
    """Create a safe filename component by sanitizing the input string.
    
    Replaces all non-word characters (excluding dots and hyphens) with underscores
    and strips leading/trailing dots and underscores to produce a valid filename
    component.
    
    Args:
        value: The input string to sanitize for use as a filename component.
        default: The fallback value to return if the sanitized result is empty.
            Defaults to "document".
    
    Returns:
        A sanitized string suitable for use in filenames, or the default value
        if sanitization results in an empty string.
    
    Example:
        >>> make_safe_filename_component("My Document (v2)")
        'My_Document__v2_'
        >>> make_safe_filename_component("")
        'document'
    """
    cleaned = SAFE_FILENAME_RE.sub("_", value).strip("._")
    return cleaned or default


def normalize_whitespace(text: str) -> str:
    """Normalize whitespace and line breaks in text content.
    
    Performs Unicode NFC normalization, converts various line break formats
    to standard newlines, and replaces special whitespace characters with
    regular spaces.
    
    Args:
        text: The input text string to normalize.
    
    Returns:
        The normalized text with consistent whitespace handling.
    
    Note:
        - Converts CRLF (\\r\\n) and CR (\\r) to LF (\\n)
        - Replaces non-breaking space (\\u00a0) and ideographic space (\\u3000)
          with regular space
        - Uses NFC (Canonical Composition) Unicode normalization
    """
    text = unicodedata.normalize("NFC", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    return text.replace("\u00a0", " ").replace("\u3000", " ")


def replace_citations(text: str) -> str:
    """Replace citation markers with styled HTML superscript links.
    
    Transforms citation patterns like [[1]](https://example.com) into
    styled HTML superscript elements with proper link formatting and
    security attributes.
    
    Args:
        text: The input text containing citation markers to transform.
    
    Returns:
        The text with citations replaced by HTML superscript links, with
        extraneous whitespace before citations removed.
    
    Note:
        - Citation format: [[number]](url)
        - Output format: <sup class="citation"><a href="url" target="_blank">[number]</a></sup>
        - URLs are HTML-escaped for security
        - Removes whitespace before <sup class="citation"> tags
    """
    def _repl(match: re.Match[str]) -> str:
        idx, url = match.group(1), match.group(2).strip()
        safe_url = html.escape(url, quote=True)
        return (
            f'<sup class="citation">'
            f'<a href="{safe_url}" target="_blank" rel="noopener noreferrer">[{idx}]</a>'
            f"</sup>"
        )

    text = CITATION_RE.sub(_repl, text)
    return re.sub(r'[ \t]+(<sup class="citation">)', r"\1", text)


def strip_internal_citation_markers(text: str) -> str:
    text = CHECKED_CITATION_RE.sub(r"\1", text)
    return LEGACY_CITATION_RE.sub("", text)


def replace_citations_with_markdown_links(text: str) -> str:
    def _repl(match: re.Match[str]) -> str:
        idx, url = match.group(1), match.group(2).strip()
        return rf"[\[{idx}\]]({url})"

    return CITATION_RE.sub(_repl, text)


def _is_pipe_table_delimiter_line(line: str) -> bool:
    stripped = line.strip()
    if not (stripped.startswith("|") and stripped.endswith("|")):
        return False

    cells = [cell.strip() for cell in stripped.strip("|").split("|")]
    if not cells:
        return False

    for cell in cells:
        if not cell:
            return False
        if re.fullmatch(r":?-{3,}:?", cell) is None:
            return False
    return True


def normalize_pipe_tables(text: str) -> str:
    lines = text.split("\n")
    result: list[str] = []
    index = 0

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if (
            stripped.startswith("|")
            and stripped.endswith("|")
            and index + 1 < len(lines)
            and _is_pipe_table_delimiter_line(lines[index + 1])
        ):
            block: list[str] = [line, lines[index + 1]]
            index += 2

            while index < len(lines):
                current = lines[index]
                current_stripped = current.strip()
                if current_stripped.startswith("|") and current_stripped.endswith("|"):
                    block.append(current)
                    index += 1
                    continue
                break

            if result and result[-1].strip():
                result.append("")
            result.extend(block)
            if index < len(lines) and lines[index].strip():
                result.append("")
            continue

        result.append(line)
        index += 1

    return "\n".join(result)


def normalize_reference_lines(text: str) -> str:
    """Normalize reference section formatting in markdown text.
    
    Detects reference/bibliography sections (marked with Chinese or English
    headings) and converts numbered reference lines (e.g., "[1]. Some text")
    to markdown list format for proper rendering.
    
    Args:
        text: The input markdown text containing a reference section.
    
    Returns:
        The text with reference lines converted to markdown list items
        within reference sections.
    
    Note:
        - Recognizes section headers: #/##/### 参考文献 or references
        - Converts "[N]. content" to "- [N] content" within reference sections
        - Preserves existing list items and blank lines
        - Exits reference section mode when encountering new headings
    """
    lines = text.splitlines()
    result: list[str] = []
    in_reference_section = False

    for line in lines:
        stripped = line.strip().lower()
        if stripped in {
            "# 参考文献",
            "## 参考文献",
            "### 参考文献",
            "# references",
            "## references",
            "### references",
        }:
            in_reference_section = True
            result.append(line)
            continue

        if in_reference_section:
            match = REFERENCE_LINE_RE.match(line)
            if match:
                indent = match.group("indent")
                idx = match.group(2)
                content = match.group(3)
                result.append(f"{indent}- [{idx}] {content}")
                continue

            if stripped == "":
                result.append(line)
                continue

            if re.match(r"^\s*[-*]\s+", line):
                result.append(line)
                continue

            if re.match(r"^\s*#{1,6}\s+", line):
                in_reference_section = False

        result.append(line)

    return "\n".join(result)


def fix_center_caption_blocks(text: str) -> str:
    """Convert center-styled divs to figure-caption divs for proper rendering.
    
    Replaces inline style-based center alignment divs with a semantic
    figure-caption class that works better with markdown processing.
    
    Args:
        text: The input text containing potentially malformed center divs.
    
    Returns:
        The text with center-styled divs replaced by figure-caption divs
        with markdown attribute enabled.
    
    Note:
        Pattern matched: <div style="text-align: center;"> (case-insensitive)
        Replacement: <div class="figure-caption" markdown="1">
    """
    return CENTER_CAPTION_RE.sub('<div class="figure-caption" markdown="1">', text)


def normalize_legacy_font_caption_blocks(text: str) -> str:
    """Rewrite legacy font-based captions into block-level caption containers."""

    def _replace(match: re.Match[str]) -> str:
        indent = match.group("indent")
        body = match.group("body").strip()
        citations = match.group("citations").strip()
        caption_content = f"{body}{citations}"
        return (
            f'\n{indent}<div class="figure-caption" markdown="1">\n'
            f"{indent}{caption_content}\n"
            f"{indent}</div>\n"
        )

    return LEGACY_FONT_CAPTION_LINE_RE.sub(_replace, text)


def normalize_interrupted_nested_list_blocks(text: str) -> str:
    """Keep image and caption blocks inside the list items they interrupt."""
    lines = text.split("\n")

    def _indent_width(indent: str) -> int:
        return len(indent.expandtabs(4))

    for index, line in enumerate(lines):
        current_item = LIST_ITEM_WITH_INDENT_RE.match(line)
        if current_item is None:
            continue

        list_indent = _indent_width(current_item.group("indent"))
        block_indexes: list[int] = []
        saw_image = False
        saw_font_caption = False

        for cursor in range(index + 1, len(lines)):
            candidate = lines[cursor]
            if not candidate.strip():
                block_indexes.append(cursor)
                continue

            if MARKDOWN_IMAGE_LINE_RE.match(candidate):
                saw_image = True
                block_indexes.append(cursor)
                continue

            if LEGACY_FONT_CAPTION_PREFIX_RE.match(candidate):
                saw_font_caption = True
                block_indexes.append(cursor)
                continue

            next_item = LIST_ITEM_WITH_INDENT_RE.match(candidate)
            next_indent = _indent_width(next_item.group("indent")) if next_item else -1
            is_matching_sibling = saw_image and next_item is not None and next_indent == list_indent
            is_nested_after_caption = (
                saw_font_caption and next_item is not None and next_indent > list_indent
            )
            if is_matching_sibling or is_nested_after_caption:
                content_indent_width = next_indent if is_nested_after_caption else list_indent + 4
                content_indent = " " * content_indent_width
                for block_index in block_indexes:
                    if lines[block_index].strip():
                        lines[block_index] = content_indent + lines[block_index].lstrip()
            break

    return "\n".join(lines)


def normalize_list_boundaries(text: str) -> str:
    """Insert a blank line before list items when the source omits one."""
    block_prefixes = ("#", ">", "|", "```", "<div", "<ul", "<ol", "<li")

    def _needs_blank_line_before_list(current_line: str, previous_line: str) -> bool:
        previous = previous_line.strip()
        if not LIST_ITEM_RE.match(current_line):
            return False
        if not previous:
            return False
        if LIST_ITEM_RE.match(previous_line):
            return False
        return not previous.startswith(block_prefixes)

    lines = text.split("\n")
    normalized: list[str] = []

    for line in lines:
        previous_line = normalized[-1] if normalized else ""
        if _needs_blank_line_before_list(line, previous_line):
            normalized.append("")
        normalized.append(line)

    return "\n".join(normalized)


def normalize_orphan_indented_list_items(text: str) -> str:
    """Dedent report-style list items that Markdown would otherwise treat as code."""
    lines = text.split("\n")
    normalized: list[str] = []
    in_fenced_code = False
    orphan_list_indent: int | None = None

    def _previous_nonempty_line() -> str:
        for previous in reversed(normalized):
            if previous.strip():
                return previous
        return ""

    def _indent_width(indent: str) -> int:
        return len(indent.expandtabs(4))

    def _line_indent_width(line: str) -> int:
        indent = re.match(r"^[ \t]*", line).group(0)
        return _indent_width(indent)

    for line in lines:
        if line.lstrip().startswith("```"):
            in_fenced_code = not in_fenced_code
            normalized.append(line)
            orphan_list_indent = None
            continue

        match = INDENTED_LIST_ITEM_RE.match(line)
        if match and not in_fenced_code:
            indent_width = _indent_width(match.group("indent"))
            if orphan_list_indent == indent_width:
                normalized.append(match.group("marker"))
                continue

            previous = _previous_nonempty_line()
            previous_indent_width = _line_indent_width(previous)
            follows_nested_content = (
                previous_indent_width > indent_width
                or (
                    previous_indent_width == indent_width
                    and INDENTED_HTML_BLOCK_END_RE.match(previous)
                )
            )
            if (
                not follows_nested_content
                and not LIST_ITEM_RE.match(previous)
                and not INDENTED_LIST_ITEM_RE.match(previous)
            ):
                orphan_list_indent = indent_width
                normalized.append(match.group("marker"))
                continue

            orphan_list_indent = None
        elif line.strip():
            orphan_list_indent = None

        normalized.append(line)

    return "\n".join(normalized)


def strip_mermaid_blocks(text: str) -> str:
    """Remove Mermaid fenced code blocks and adjacent figure captions.

    Mermaid charts are now generated upstream as regular image references.
    Any leftover Mermaid fences are treated as stale intermediate content and
    removed before HTML or DOCX conversion. An immediately following caption
    block is removed as well so no orphaned figure note remains.
    """
    return MERMAID_BLOCK_WITH_ADJACENT_CAPTION_RE.sub("\n", text)


def preprocess_markdown_text(text: str) -> str:
    """Apply preprocessing transformations to markdown text.
    
    Chains multiple normalization functions to prepare markdown content
    for conversion, handling whitespace, citations, references, and
    caption blocks.
    
    Args:
        text: The raw markdown text to preprocess.
    
    Returns:
        The preprocessed text ready for conversion operations.
    
    Note:
        Processing order: whitespace normalization → citation replacement
        → reference line normalization → center caption block fixing.
    """
    text = normalize_whitespace(text)
    text = normalize_interrupted_nested_list_blocks(text)
    text = strip_internal_citation_markers(text)
    text = replace_citations(text)
    text = normalize_pipe_tables(text)
    text = normalize_reference_lines(text)
    text = fix_center_caption_blocks(text)
    text = normalize_legacy_font_caption_blocks(text)
    text = normalize_orphan_indented_list_items(text)
    text = normalize_list_boundaries(text)
    return strip_mermaid_blocks(text)


def _filter_remote_urls_for_ssrf(text: str) -> str:
    """安全修复: 过滤 Markdown 中的远程 URL 以防止 SSRF.
    
    Pandoc 在转换时会拉取 Markdown 中引用的远程图片/资源,
    这可能导致 SSRF 攻击。此函数移除或替换远程 URL 引用。
    
    Args:
        text: 包含潜在远程 URL 的 Markdown 文本
        
    Returns:
        过滤后的 Markdown 文本,远程 URL 被移除或替换为占位符
    """
    # 匹配 Markdown 图片引用: ![alt](http://...) 或 ![alt](https://...)
    # 保留本地相对路径和数据 URI,移除远程 http/https URL
    remote_image_pattern = re.compile(
        r'!\[([^\]]*)\]\((https?://[^\s)]+)\)',
        re.IGNORECASE
    )
    
    # 匹配 Markdown 链接引用: [text](http://...) 或 [text](https://...)
    # 注意: 不要移除普通链接,只对资源拉取相关的进行处理
    # Pandoc 主要拉取图片资源,所以这里主要处理图片
    
    def replace_remote_image(match):
        alt_text = match.group(1)
        url = match.group(2)
        # 安全替换: 使用占位符替代远程图片
        logger.warning(
            "[Security] Remote image URL filtered for SSRF prevention: %s",
            url
        )
        return f"[远程图片已过滤: {alt_text}]"
    
    return remote_image_pattern.sub(replace_remote_image, text)


def preprocess_markdown_text_for_docx(text: str) -> str:
    text = normalize_whitespace(text)
    text = normalize_interrupted_nested_list_blocks(text)
    text = strip_internal_citation_markers(text)
    text = replace_citations(text)
    text = normalize_pipe_tables(text)
    text = normalize_reference_lines(text)
    text = _filter_remote_urls_for_ssrf(text)
    text = fix_center_caption_blocks(text)
    text = normalize_legacy_font_caption_blocks(text)
    text = normalize_orphan_indented_list_items(text)
    text = normalize_list_boundaries(text)
    return strip_mermaid_blocks(text)


def postprocess_html(html_text: str) -> str:
    """Apply postprocessing transformations to converted HTML output.
    
    Enhances HTML by adding security attributes to external links and
    wrapping citation anchors in superscript elements for consistent
    styling.
    
    Args:
        html_text: The HTML text output from conversion to postprocess.
    
    Returns:
        The postprocessed HTML with enhanced links and citation styling.
    
    Note:
        - External links get target="_blank" and rel="noopener noreferrer"
        - Citation anchors ([N]) are wrapped in <sup class="citation">
        - Whitespace around citation elements is cleaned
    """
    def _repl(match: re.Match[str]) -> str:
        before = match.group(1).rstrip()
        href = re.sub(r"\s+", "", match.group(2))
        after = match.group(3).rstrip()
        attrs = " ".join(part for part in [before, f'href="{href}"', after] if part)
        return f'<a {attrs} target="_blank" rel="noopener noreferrer">'

    def _wrap_citation(match: re.Match[str]) -> str:
        return f'<sup class="citation">{match.group(1)}</sup>'

    html_text = EXTERNAL_LINK_RE.sub(_repl, html_text)
    html_text = CITATION_ANCHOR_RE.sub(_wrap_citation, html_text)
    html_text = re.sub(r'[ \t]+(<sup class="citation">)', r"\1", html_text)
    return re.sub(r'(</sup>)[ \t]+(<sup class="citation">)', r"\1\2", html_text)


def protect_math_spans(text: str) -> tuple[str, list[str]]:
    """Protect LaTeX math spans from Markdown emphasis/link parsing."""
    formulas: list[str] = []
    code_spans: list[str] = []

    def _store_code(match: re.Match[str]) -> str:
        code_spans.append(match.group(0))
        return CODE_PLACEHOLDER.format(len(code_spans) - 1)

    def _store_math(match: re.Match[str]) -> str:
        formulas.append(match.group(0))
        return MATH_PLACEHOLDER.format(len(formulas) - 1)

    text = FENCED_CODE_RE.sub(_store_code, text)
    text = INLINE_CODE_RE.sub(_store_code, text)
    text = BLOCK_MATH_RE.sub(_store_math, text)
    text = _protect_inline_math_spans(text, formulas)

    def _restore_code(match: re.Match[str]) -> str:
        index = int(match.group(1))
        if 0 <= index < len(code_spans):
            return code_spans[index]
        return match.group(0)

    return CODE_PLACEHOLDER_RE.sub(_restore_code, text), formulas


def _protect_inline_math_spans(text: str, formulas: list[str]) -> str:
    parts: list[str] = []
    cursor = 0
    index = 0

    while index < len(text):
        start = text.find("$", index)
        if start == -1:
            break
        if _should_skip_inline_math_start(text, start):
            index = start + 1
            continue

        end = _find_inline_math_end(text, start + 1)
        if end is None:
            index = start + 1
            continue

        formula = text[start:end + 1]
        if _is_likely_inline_math(formula[1:-1].strip()):
            parts.append(text[cursor:start])
            formulas.append(formula)
            parts.append(MATH_PLACEHOLDER.format(len(formulas) - 1))
            cursor = end + 1
            index = end + 1
        else:
            index = end + 1

    if not parts:
        return text

    parts.append(text[cursor:])
    return "".join(parts)


def _find_inline_math_end(text: str, start_index: int) -> int | None:
    index = start_index
    while index < len(text):
        end = text.find("$", index)
        if end == -1:
            return None
        if "\n" in text[start_index:end]:
            return None
        if _is_valid_inline_math_end(text, start_index, end):
            return end
        index = end + 1
    return None


def _should_skip_inline_math_start(text: str, index: int) -> bool:
    if _is_escaped(text, index):
        return True
    if _is_double_dollar(text, index):
        return True
    if _is_likely_currency_start(text, index):
        return True
    return not _is_valid_inline_math_start(text, index)


def _is_valid_inline_math_end(text: str, start_index: int, end_index: int) -> bool:
    if _is_escaped(text, end_index):
        return False
    if _is_double_dollar(text, end_index):
        return False
    if end_index <= start_index:
        return False
    if text[end_index - 1].isspace():
        return False
    return not (end_index + 1 < len(text) and text[end_index + 1].isdigit())


def _is_likely_inline_math(content: str) -> bool:
    if not content:
        return False
    if HTML_TAG_RE.search(content) or MARKDOWN_LINK_RE.search(content):
        return False
    if PLAIN_CURRENCY_RE.fullmatch(content):
        return False
    return bool(
        SIMPLE_VARIABLE_MATH_RE.fullmatch(content)
        or COMPACT_VARIABLE_MATH_RE.fullmatch(content)
        or ALNUM_VARIABLE_MATH_RE.fullmatch(content)
        or VARIABLE_LIST_MATH_RE.fullmatch(content)
        or PAREN_VARIABLE_LIST_MATH_RE.fullmatch(content)
        or MATH_FEATURE_RE.search(content)
        or MATH_FUNCTION_CALL_RE.fullmatch(content)
        or PRIME_FUNCTION_CALL_RE.fullmatch(content)
        or FUNCTION_CALL_SEQUENCE_RE.fullmatch(content)
        or ABSOLUTE_VALUE_MATH_RE.fullmatch(content)
        or NORM_MATH_RE.fullmatch(content)
        or PERCENT_MATH_RE.fullmatch(content)
        or _is_balanced_math_function_call(content)
        or EQUATION_REFERENCE_RE.fullmatch(content)
        or NUMERIC_TUPLE_RE.fullmatch(content)
        or PRIME_VARIABLE_RE.fullmatch(content)
    )


def _is_balanced_math_function_call(content: str) -> bool:
    name_match = re.match(
        r"(?:\\?[A-Za-z\u0370-\u03ff]+|[0-9]+|\\[A-Za-z]+)"
        r"\s*(?:_\{?[\w\\]+\}?|\^\{?[\w\\]+\}?|\\[A-Za-z]+)*\s*",
        content,
    )
    if name_match is None:
        return False

    open_index = name_match.end()
    if open_index >= len(content) or content[open_index] not in "([":
        return False

    pairs = {"(": ")", "[": "]"}
    closing_char = pairs.get(content[open_index])
    if closing_char is None:
        return False
    stack = [closing_char]
    index = open_index + 1
    saw_argument = False

    while index < len(content):
        char = content[index]
        if char == "\\":
            index += 2
            saw_argument = True
            continue
        closing_char = pairs.get(char)
        if closing_char is not None:
            stack.append(closing_char)
            saw_argument = True
        elif stack and char == stack[-1]:
            stack.pop()
            if not stack:
                return saw_argument and index == len(content) - 1
        elif char in ")]":
            return False
        elif not char.isspace():
            saw_argument = True
        index += 1

    return False


def _is_valid_inline_math_start(text: str, index: int) -> bool:
    next_index = index + 1
    if next_index >= len(text):
        return False
    next_char = text[next_index]
    if next_char.isspace() or next_char in "，。；：、.!?;:)]}）】」』":
        return False
    return True


def _is_likely_currency_start(text: str, index: int) -> bool:
    match = re.match(r"\$(?:\d{1,3}(?:,\d{3})*|\d+)(?:\.\d+)?", text[index:])
    if not match:
        return False
    next_index = index + len(match.group(0))
    if next_index >= len(text):
        return True
    if text.startswith(r"\%", next_index) or text[next_index] == "%":
        return False
    end = _find_inline_math_end(text, index + 1)
    if end is not None and "|" in text[next_index:end]:
        return True
    if end is not None and STRONG_NUMERIC_MATH_FEATURE_RE.search(text[next_index:end]):
        return False
    return text[next_index] not in "([{_^=\\"


def _is_double_dollar(text: str, index: int) -> bool:
    return (
        (index > 0 and text[index - 1] == "$")
        or (index + 1 < len(text) and text[index + 1] == "$")
    )


def _is_escaped(text: str, index: int) -> bool:
    slash_count = 0
    cursor = index - 1
    while cursor >= 0 and text[cursor] == "\\":
        slash_count += 1
        cursor -= 1
    return slash_count % 2 == 1


def restore_math_spans(html_text: str, formulas: list[str]) -> str:
    """Restore math placeholders produced by protect_math_spans."""

    def _restore(match: re.Match[str]) -> str:
        index = int(match.group(1))
        if 0 <= index < len(formulas):
            return normalize_math_span_for_rendering(formulas[index])
        return match.group(0)

    return MATH_PLACEHOLDER_RE.sub(_restore, html_text)


def normalize_math_span_for_rendering(formula: str) -> str:
    """Normalize protected dollar math to MathJax-friendly delimiters."""
    match = INLINE_MATH_DOLLAR_RE.match(formula)
    if match:
        return rf"\({html.escape(match.group(1), quote=False)}\)"
    if BLOCK_MATH_RE.fullmatch(formula):
        return f"$${html.escape(formula[2:-2], quote=False)}$$"
    return html.escape(formula, quote=False)


def _neighbor_numbered_line(lines: list[str], index: int, *, reverse: bool) -> str | None:
    """Find the nearest non-empty line in a specified direction.
    
    Private helper function that searches forward or backward from a given
    index to find the first non-empty, non-whitespace line.
    
    Args:
        lines: List of text lines to search through.
        index: The starting index position in the lines list.
        reverse: If True, search backward; if False, search forward.
    
    Returns:
        The first non-empty line found in the specified direction, or None
        if no such line exists before reaching the list boundary.
    
    Note:
        Empty lines (stripped content is empty) are skipped during search.
    """
    step = -1 if reverse else 1
    cursor = index + step

    while 0 <= cursor < len(lines):
        stripped = lines[cursor].strip()
        if stripped:
            return stripped
        cursor += step

    return None


def _should_promote_numbered_heading(lines: list[str], index: int, match: re.Match[str]) -> bool:
    """Determine if a numbered line should be promoted to a markdown heading.
    
    Private helper function that applies heuristics to decide whether
    a numbered pattern (e.g., "1. Introduction") represents a heading
    rather than a regular list item.
    
    Args:
        lines: The full list of text lines for context analysis.
        index: The index of the current line being evaluated.
        match: The regex match object for the numbered heading pattern.
    
    Returns:
        True if the line should be converted to a markdown heading,
        False if it should remain as a regular numbered line.
    
    Note:
        Promotion criteria:
        - Title length must be <= 80 characters
        - Title must not end with sentence-ending punctuation
        - Line must be surrounded by blank lines (no content immediately before/after)
        - Adjacent non-empty lines must not also match numbered heading pattern
    """
    title = match.group("title").strip()

    if len(title) > 80 or SENTENCE_END_RE.search(title):
        return False

    if index > 0 and lines[index - 1].strip():
        return False

    if index + 1 < len(lines) and lines[index + 1].strip():
        return False

    prev_nonempty = _neighbor_numbered_line(lines, index, reverse=True)
    next_nonempty = _neighbor_numbered_line(lines, index, reverse=False)
    if prev_nonempty and NUMBERED_HEADING_RE.match(prev_nonempty):
        return False
    if next_nonempty and NUMBERED_HEADING_RE.match(next_nonempty):
        return False

    return True


def normalize_headings(content: str) -> str:
    """Normalize heading structure in markdown content.
    
    Processes markdown text to ensure proper heading formatting with
    appropriate spacing. Converts detected numbered patterns (e.g., "1.2 Title")
    to markdown headings when they meet promotion criteria, and ensures
    standard hash-based headings are properly spaced.
    
    Args:
        content: The raw markdown text content to normalize.
    
    Returns:
        The normalized markdown with properly formatted headings, blank
        lines between sections, and excess newlines collapsed.
    
    Note:
        - Preserves code blocks without modification
        - Ensures blank lines before and after headings
        - Converts numbered patterns to headings based on heuristics
        - Limits heading level to 6 (maximum markdown level)
        - Collapses consecutive blank lines to maximum of two
    """
    content = normalize_whitespace(content)

    lines = content.split("\n")
    out: list[str] = []
    in_code_block = False

    for index, raw_line in enumerate(lines):
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code_block = not in_code_block
            out.append(line)
            continue

        if in_code_block:
            out.append(line)
            continue

        if not stripped:
            out.append("")
            continue

        match_hash = re.match(r"^(#{1,6})\s*(.+?)\s*$", stripped)
        if match_hash:
            hashes = match_hash.group(1)
            title = match_hash.group(2).strip()
            if out and out[-1] != "":
                out.append("")
            out.append(f"{hashes} {title}")
            out.append("")
            continue

        match_numbered = NUMBERED_HEADING_RE.match(line)
        if match_numbered and _should_promote_numbered_heading(lines, index, match_numbered):
            numbering = match_numbered.group("number")
            title = match_numbered.group("title").strip()
            level = min(numbering.count(".") + 1, 6)
            heading = "#" * level + " " + f"{numbering} {title}"
            if out and out[-1] != "":
                out.append("")
            out.append(heading)
            out.append("")
            continue

        out.append(line)

    result = "\n".join(out)
    result = re.sub(r"\n{3,}", "\n\n", result)
    return result.strip() + "\n"


def _set_rfonts(rpr, font_name: str) -> None:
    """Set font properties on an rPr (run properties) XML element.
    
    Private helper function that creates or updates the rFonts element
    within run properties, setting all font categories to the specified
    font name for comprehensive font coverage.
    
    Args:
        rpr: The OxmlElement representing run properties (w:rPr).
        font_name: The font name to apply to all font categories.
    
    Returns:
        None
    
    Note:
        Sets four font categories: ascii, hAnsi, eastAsia, and cs
        to ensure consistent rendering across different character sets.
    """
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)

    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), font_name)


def _set_run_font(run, font_name: str) -> None:
    """Set font on a docx text run object.
    
    Private helper function that applies a font name to both the
    high-level run.font property and the underlying XML rFonts element.
    
    Args:
        run: The docx Run object to modify.
        font_name: The font name to apply.
    
    Returns:
        None
    
    Note:
        Uses get_or_add_rPr() to ensure rPr element exists before setting fonts.
    """
    run.font.name = font_name
    _set_rfonts(run.element.get_or_add_rPr(), font_name)


def _set_style_font(style, font_name: str) -> None:
    """Set font on a docx style object.
    
    Private helper function that applies a font name to both the
    high-level style.font property and the underlying XML rFonts element.
    
    Args:
        style: The docx Style object to modify.
        font_name: The font name to apply.
    
    Returns:
        None
    
    Note:
        Uses get_or_add_rPr() to ensure rPr element exists before setting fonts.
    """
    style.font.name = font_name
    _set_rfonts(style.element.get_or_add_rPr(), font_name)


def _apply_font_to_table(table, font_name: str) -> None:
    """Apply font to all text content within a docx table.
    
    Private helper function that recursively processes all cells,
    paragraphs, and runs in a table, including nested tables,
    to apply uniform font styling.
    
    Args:
        table: The docx Table object to process.
        font_name: The font name to apply to all text content.
    
    Returns:
        None
    
    Note:
        Recursively handles nested tables within cells for complete coverage.
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run, font_name)
            for nested_table in cell.tables:
                _apply_font_to_table(nested_table, font_name)


def normalize_docx_fonts(docx_path: Path, *, font_name: str = DEFAULT_DOCX_FONT) -> None:
    """Normalize fonts throughout a Word document for consistent rendering.
    
    Applies the specified font to all text elements in the document including
    styles, paragraphs, tables, headers, and footers. Ensures uniform font
    appearance across all document components.
    
    Args:
        docx_path: Path to the .docx file to process. The file is modified in-place.
        font_name: The font name to apply throughout the document.
            Defaults to DEFAULT_DOCX_FONT ("Microsoft YaHei").
    
    Returns:
        None. The document is modified and saved in-place.
    
    Note:
        - Requires python-docx library (DOCX_STYLE_AVAILABLE flag)
        - Logs a warning and returns early if python-docx is unavailable
        - Processes standard styles (Normal, Title, Heading 1-9, etc.)
        - Handles tables, headers, footers, and nested content
        - Missing styles are silently skipped (KeyError handling)
    """
    if not DOCX_STYLE_AVAILABLE:
        logger.warning("python-docx is unavailable, skipping font normalization for %s.", docx_path)
        return

    document = Document(docx_path)

    for style_name in (
        "Normal",
        "Title",
        "Subtitle",
        "Body Text",
        "Hyperlink",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Heading 4",
        "Heading 5",
        "Heading 6",
        "Heading 7",
        "Heading 8",
        "Heading 9",
    ):
        try:
            _set_style_font(document.styles[style_name], font_name)
        except KeyError:
            continue

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            _set_run_font(run, font_name)

    for table in document.tables:
        _apply_font_to_table(table, font_name)

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            for run in paragraph.runs:
                _set_run_font(run, font_name)
        for paragraph in section.footer.paragraphs:
            for run in paragraph.runs:
                _set_run_font(run, font_name)
        for table in section.header.tables:
            _apply_font_to_table(table, font_name)
        for table in section.footer.tables:
            _apply_font_to_table(table, font_name)

    document.save(docx_path)


def normalize_docx_tables(docx_path: Path) -> None:
    """Apply simple, readable defaults to generated DOCX tables."""
    if not DOCX_STYLE_AVAILABLE:
        logger.warning("python-docx is unavailable, skipping table normalization for %s.", docx_path)
        return

    document = Document(docx_path)
    changed = False
    for table in document.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        changed = True

    if changed:
        document.save(docx_path)


def enhance_image(image_path: str) -> None:
    """Enhance image quality for document embedding.
    
    Upscales the image, applies contrast, sharpness, and color enhancements,
    and saves as optimized PNG format. Handles transparency by compositing
    onto a white background.
    
    Args:
        image_path: Path to the image file to enhance. The file is modified
            in-place and saved as PNG format.
    
    Returns:
        None. The image file is modified in-place.
    
    Note:
        - Requires PIL/Pillow library (PIL_AVAILABLE flag)
        - Returns silently if PIL is unavailable
        - Upscale factor: DOCX_IMAGE_UPSCALE_FACTOR (3x)
        - Enhancement factors: contrast, sharpness, color from constants
        - RGBA/LA images are composited onto white background before conversion
        - On failure, logs a warning but does not raise exception
    """
    if not PIL_AVAILABLE:
        return

    try:
        with Image.open(image_path) as original:
            if original.mode in ("RGBA", "LA"):
                background = Image.new("RGBA", original.size, (255, 255, 255, 255))
                background.alpha_composite(original.convert("RGBA"))
                image = background.convert("RGB")
            else:
                image = original.convert("RGB")

        image = image.resize(
            (image.width * DOCX_IMAGE_UPSCALE_FACTOR, image.height * DOCX_IMAGE_UPSCALE_FACTOR),
            Image.Resampling.LANCZOS,
        )
        image = ImageEnhance.Contrast(image).enhance(DOCX_IMAGE_CONTRAST)
        image = ImageEnhance.Sharpness(image).enhance(DOCX_IMAGE_SHARPNESS)
        image = ImageEnhance.Color(image).enhance(DOCX_IMAGE_COLOR)
        image.save(image_path, format="PNG", optimize=True)
    except Exception as exc:
        logger.warning("Failed to enhance image %s: %s", image_path, exc)
