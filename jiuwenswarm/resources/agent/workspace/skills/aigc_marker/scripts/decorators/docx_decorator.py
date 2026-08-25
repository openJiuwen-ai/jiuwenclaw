"""DOCX AIGC decorator."""
import os.path
import zipfile
from xml.etree import ElementTree as ET

import docx
from docx.shared import Pt
from docx_extend.api import DocumentExtend

from decorators.common import DEFAULT_CUSTOM_PROPERTY_FMTID, get_aigc_signature, is_aigc_complete, parse_aigc_json


class DocxAigcDecorator:
    """DOCX AIGC decorator - adds hidden AIGC mark to Word documents."""

    def __init__(self):
        self.name = "docx_aigc_decorator"

    def decorate(self, file_path: str, content: str, add_visible_mark: bool = True):
        """Add AIGC mark to DOCX file."""
        existing = self._get_aigc_data(file_path)
        if existing and is_aigc_complete(existing):
            print(f"  [DOCX] AIGC mark complete, skipping")
            return
        try:
            aigc_signature = get_aigc_signature(content)
            self._add_aigc_mark(file_path, aigc_signature, add_visible_mark)
            print(f"  [DOCX] AIGC mark added successfully")
        except Exception as e:
            print(f"  [DOCX] Warning: Failed to add AIGC mark: {str(e)}")

    def _get_aigc_data(self, file_path: str) -> dict | None:
        """Read and parse AIGC metadata from DOCX custom properties."""
        try:
            with zipfile.ZipFile(file_path, 'r') as zf:
                if 'docProps/custom.xml' not in zf.namelist():
                    return None
                xml_content = zf.read('docProps/custom.xml')
                root = ET.fromstring(xml_content)
                for prop in root.iter():
                    if prop.get('name') == 'AIGC':
                        for child in prop:
                            if child.text:
                                return parse_aigc_json(child.text)
                        return None
        except (zipfile.BadZipFile, ET.ParseError, KeyError):
            return None
        return None

    def _add_aigc_mark(self, file_path: str, signature: str, add_visible_mark: bool = True):
        """Add custom property to DOCX file using docx library."""
        if not file_path or not os.path.exists(file_path):
            raise FileNotFoundError(f"Target file not found: {file_path}")

        doc = docx.Document(file_path)
        if add_visible_mark:
            self._add_visible_mark_to_docx(doc)
        self._add_implicit_mark(doc, signature)
        doc.save(file_path)

    def _add_visible_mark_to_docx(self, doc: docx.Document):
        """Add visible AIGC mark paragraph to DOCX file."""

        # 添加段落
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("内容由AI生成")

        # 设置字体为宋体，5号字(10.5pt)
        run.font.name = "宋体"
        run.font.size = Pt(10.5)

    def _add_implicit_mark(self, doc, signature):
        doc_ex = DocumentExtend(doc)
        custom_properties_part = doc_ex.custom_properties_part
        custom_properties = custom_properties_part.custom_properties

        # Check if AIGC property already exists and update its value
        for prop in custom_properties._element:
            if prop.get('name') == 'AIGC':
                for child in prop:
                    child.text = signature
                return

        # Not found, add new property
        pid = custom_properties_part.next_id
        custom_properties.add_property("AIGC", signature, DEFAULT_CUSTOM_PROPERTY_FMTID, pid)
