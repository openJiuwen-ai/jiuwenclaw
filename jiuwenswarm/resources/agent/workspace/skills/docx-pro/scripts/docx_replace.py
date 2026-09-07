# -*- coding: utf-8 -*-
"""
docx_replace.py - Word (.docx) 保真文本替换引擎（docx-pro 技能）

解决 Word 文本替换的两大痛点：

1. 跨 run 文本定位
   Word 常因拼写检查、协同编辑、格式刷等原因把一段连续文字拆成多个 run
   （w:r/w:t 节点）。对 document.xml 做朴素字符串替换会大量漏换。
   本引擎把同一段落内相邻 w:t 的文本拼接后定位目标串：替换文本写入首个
   命中节点，其余被覆盖节点的文本清空——不增删任何 XML 元素，run 结构与
   全部格式属性（字体/颜色/字号/样式）零改动。

2. 字节级保真
   zip 包内仅含目标文本的 XML 部件被修改；其余条目（样式表、主题、编号、
   字体表、图片……）按原始字节复制。命中部件也只改动被替换节点的文本内容，
   其余 XML 原样保留（保留原 XML 声明、属性顺序、命名空间前缀）。

覆盖范围：正文（含表格单元格、文本框、超链接内文本）、页眉、页脚、
脚注、尾注、批注（--scope all）；或仅正文（--scope body）。

安全边界：匹配不跨段落、不跨表格单元格（相邻 w:t 之间出现段落结束、
制表符、换行符、域代码等即视为文本不连续）；域代码 w:instrText 与
修订删除文本 w:delText 不参与匹配。

依赖：仅 Python 标准库（zipfile / re / json / bisect / difflib）。
"""

import bisect
import json
import re
import zipfile

__all__ = [
    "replace_in_docx", "replace_in_xml", "count_in_docx", "REPLACE_ALL_PARTS",
]

SENTINEL = "\x00"  # 段落/表格/制表符等造成的文本断点标记，永不参与匹配

# 匹配 <w:t ...>text</w:t>（不会误配 <w:tab/>、<w:tc> 等：其后必须紧跟空白或 >）
WT_RE = re.compile(r"<w:t(\s[^>]*)?>(.*?)</w:t>", re.S)

# 相邻 w:t 之间出现以下任一内容时，文本视为不连续（插入断点，禁止跨它们匹配）
BREAK_PAT = re.compile(
    r"</w:p"          # 段落结束（同时覆盖表格单元格/文本框边界）
    r"|<w:tab"        # 制表符
    r"|<w:br"         # 换行
    r"|<w:cr"         # 回车
    r"|<w:instrText"  # 域代码
    r"|<w:softHyphen"
    r"|<w:noBreakHyphen"
    r"|<w:sym"        # 符号字符
)

_ENTITY_MAP = {"amp": "&", "lt": "<", "gt": ">", "quot": '"', "apos": "'"}
_ENTITY_RE = re.compile(r"&(#[0-9]+|#x[0-9a-fA-F]+|amp|lt|gt|quot|apos);")

# 参与替换的部件（--scope all）
REPLACE_ALL_PARTS = (
    "word/document.xml", "word/footnotes.xml", "word/endnotes.xml",
)
REPLACE_ALL_PREFIXES = ("word/header", "word/footer", "word/comments")
REPLACE_BODY_PARTS = ("word/document.xml",)


# ---------------------------------------------------------------- XML 实体

def _unescape(s):
    """XML 文本节点内容 -> 原始文本"""
    def _sub(m):
        e = m.group(1)
        if e.startswith("#"):
            try:
                return chr(int(e[2:], 16) if e[1] in "xX" else int(e[1:]))
            except (ValueError, OverflowError):
                return m.group(0)
        return _ENTITY_MAP.get(e, m.group(0))
    return _ENTITY_RE.sub(_sub, s)


def _escape(s):
    """原始文本 -> XML 文本节点内容（最小转义，合法且保持可读）"""
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


# ---------------------------------------------------------------- 部件分析

class _PartAnalysis(object):
    """对一个 XML 部件的分析结果：w:t 节点序列 + 拼接文本 + 位置映射"""

    def __init__(self, xml_text):
        self.xml = xml_text
        self.matches = []       # WT_RE.match 对象列表（原文位置）
        self.inner = []         # 各节点反转义后的原始文本
        self.segments = []      # [(拼接文本起点, 节点序号, 文本)]
        self.concat = []        # 逐段拼接（断点处插入 SENTINEL）
        pos = 0                 # 拼接文本中的字符位置（注意：不是列表元素个数！）
        for m in WT_RE.finditer(xml_text):
            self.matches.append(m)
            self.inner.append(_unescape(m.group(2)))

        prev_end = None
        for i, m in enumerate(self.matches):
            if prev_end is not None:
                gap = xml_text[prev_end:m.start()]
                if BREAK_PAT.search(gap):
                    self.concat.append(SENTINEL)  # 断点：段落/单元格/制表符/域代码边界
                    pos += 1
            self.segments.append((pos, i, self.inner[i]))
            self.concat.append(self.inner[i])
            pos += len(self.inner[i])
            prev_end = m.end()
        self.C = "".join(self.concat)
        self._starts = [s[0] for s in self.segments]

    def locate(self, pos):
        """拼接文本位置 -> (节点序号, 节点内偏移)。断点位置不可定位。"""
        i = bisect.bisect_right(self._starts, pos) - 1
        if i < 0:
            raise IndexError("位置 %d 不在任何文本节点内" % pos)
        cstart, idx, text = self.segments[i]
        off = pos - cstart
        if off < 0 or off >= len(text):
            raise IndexError("位置 %d 位于文本断点上" % pos)
        return idx, off


# ---------------------------------------------------------------- 出现查找

def _find_occurrences(concat, old, case_sensitive=True):
    """在拼接文本中查找 old 的全部非重叠出现（跳过断点由 SENTINEL 天然保证）"""
    if not old:
        return []
    if SENTINEL in old:
        return []
    hay = concat if case_sensitive else concat.casefold()
    needle = old if case_sensitive else old.casefold()
    out = []
    start = 0
    while True:
        i = hay.find(needle, start)
        if i < 0:
            return out
        out.append((i, i + len(old)))
        start = i + len(old)


# ---------------------------------------------------------------- 单部件替换

def replace_in_xml(xml_text, old, new, case_sensitive=True, count_only=False):
    """对一个 XML 部件执行一次查找（及替换）。

    返回 (新xml或None, 统计dict)：
      stats = {"found": 出现次数, "spans": 涉及节点数, "parts_spanned": 跨节点出现数}
    count_only=True 时只统计，不构造新 XML。
    """
    pa = _PartAnalysis(xml_text)
    occ = _find_occurrences(pa.C, old, case_sensitive)
    stats = {"found": len(occ), "spans": 0, "cross_run": 0}
    if not occ or count_only:
        return (None, stats)

    # 收集每个节点的编辑区间（基于节点原始文本坐标，互不重叠）
    edits = {}  # 节点序号 -> [(start, end, 替换文本), ...]
    touched = set()
    for cs, ce in occ:
        s1, o1 = pa.locate(cs)
        s2, o2 = pa.locate(ce - 1)
        touched.update(range(s1, s2 + 1))
        if s1 == s2:
            edits.setdefault(s1, []).append((o1, o2 + 1, new))
        else:
            stats["cross_run"] += 1
            edits.setdefault(s1, []).append((o1, len(pa.inner[s1]), new))
            for k in range(s1 + 1, s2):
                edits.setdefault(k, []).append((0, len(pa.inner[k]), ""))
            edits.setdefault(s2, []).append((0, o2 + 1, ""))
    stats["spans"] = len(touched)

    # 逐节点应用编辑，生成新节点文本
    new_inner = dict()
    for idx, es in edits.items():
        es.sort(key=lambda e: e[0])
        pieces = []
        cur = 0
        for s, e, txt in es:
            if s < cur:
                raise RuntimeError("内部错误：节点 %d 编辑区间重叠" % idx)
            pieces.append(pa.inner[idx][cur:s])
            pieces.append(txt)
            cur = e
        pieces.append(pa.inner[idx][cur:])
        new_inner[idx] = "".join(pieces)

    # 从后往前回写原文（各节点区间互不相交，位置不会失效）
    out = xml_text
    for idx in sorted(new_inner.keys(), reverse=True):
        m = pa.matches[idx]
        text = new_inner[idx]
        open_tag = m.group(0)[: m.start(2) - m.start()]
        # 前导/尾随空白且缺少 xml:space 时补上，防止 Word 丢空格
        if text and (text[0].isspace() or text[-1].isspace()) \
                and "xml:space" not in open_tag:
            open_tag = open_tag[:-1] + ' xml:space="preserve">'
        replacement = open_tag + _escape(text) + "</w:t>"
        out = out[: m.start()] + replacement + out[m.end():]
    return (out, stats)


# ---------------------------------------------------------------- 整包替换

def _target_parts(names, scope):
    if scope == "body":
        return [n for n in names if n in REPLACE_BODY_PARTS]
    parts = [n for n in names if n in REPLACE_ALL_PARTS]
    parts += [n for n in names
              if n.endswith(".xml") and n.startswith(REPLACE_ALL_PREFIXES)]
    return parts


def replace_in_docx(src, dst, pairs, scope="all", case_sensitive=True,
                    dry_run=False, verify=True):
    """对整个 docx 执行（可能多对的）保真替换。

    pairs: [(old, new), ...] —— 按顺序链式执行（前一对的结果是后一对的输入）
    返回报告 dict；dry_run=True 时不写出文件。
    """
    if not pairs:
        raise ValueError("pairs 为空")
    with zipfile.ZipFile(src) as z:
        names = z.namelist()
        targets = _target_parts(names, scope)
        contents = {n: z.read(n) for n in names}
        infos = z.infolist()

    pair_stats = []      # [(old, new, {部件: found}), ...]
    modified = {}        # 部件名 -> 新文本
    for old, new in pairs:
        per_part = {}
        for part in targets:
            xml = modified.get(part)
            if xml is None:
                xml = contents[part].decode("utf-8")
            new_xml, st = replace_in_xml(xml, old, new, case_sensitive,
                                         count_only=dry_run)
            per_part[part] = st
            if not dry_run and new_xml is not None:
                modified[part] = new_xml
        pair_stats.append((old, new, per_part))

    report = {
        "src": src, "dst": None if dry_run else dst, "scope": scope,
        "dry_run": dry_run, "parts_scanned": targets,
        "pairs": pair_stats, "total_found": sum(
            st["found"] for _, _, pp in pair_stats for st in pp.values()),
    }
    if dry_run:
        return report

    # 重打包：仅 modified 部件换内容，其余条目按原字节/原压缩方式复制
    with zipfile.ZipFile(dst, "w") as zout:
        for item in infos:
            if item.filename in modified:
                data = modified[item.filename].encode("utf-8")
            else:
                data = contents[item.filename]
            zi = zipfile.ZipInfo(item.filename, date_time=item.date_time)
            zi.compress_type = item.compress_type
            zi.external_attr = item.external_attr
            zout.writestr(zi, data)
    report["parts_changed"] = sorted(modified.keys())

    if verify:
        report.update(_verify(src, dst, modified, pairs, targets,
                              case_sensitive))
    return report


def _verify(src, dst, modified, pairs, targets, case_sensitive):
    """替换后验证：条目一致性 / 字节级差异 / 残留计数 / 可打开性"""
    import os

    v = {}
    with zipfile.ZipFile(src) as z1, zipfile.ZipFile(dst) as z2:
        n1, n2 = z1.namelist(), z2.namelist()
        v["entry_list_identical"] = n1 == n2
        changed, unexpected = [], []
        for name in n1:
            if z1.read(name) != z2.read(name):
                if name in modified:
                    changed.append(name)
                else:
                    unexpected.append(name)
        v["changed_entries"] = sorted(changed)
        v["unexpected_changes"] = sorted(unexpected)
        v["other_entries_byte_identical"] = not unexpected

        # 残留计数（新文本若自身包含旧文本，残留>0 属正常，仅提示）
        residual = {}
        for old, new in pairs:
            cnt = 0
            for name in targets:
                xml = z2.read(name).decode("utf-8")
                _, st = replace_in_xml(xml, old, new, case_sensitive,
                                       count_only=True)
                cnt += st["found"]
            residual[old] = cnt
        v["residual"] = residual

    # python-docx 复检
    try:
        from docx import Document
        d1, d2 = Document(src), Document(dst)
        v["opens_ok"] = True
        v["paragraphs"] = (len(d1.paragraphs), len(d2.paragraphs))
        v["tables"] = (len(d1.tables), len(d2.tables))
    except ImportError:
        v["opens_ok"] = None
    except Exception as e:
        v["opens_ok"] = False
        v["open_error"] = str(e)
    v["size"] = (os.path.getsize(src), os.path.getsize(dst))
    return v


# ---------------------------------------------------------------- 计数工具

def count_in_docx(path, old, scope="all", case_sensitive=True):
    """只统计 old 在 docx 中的出现次数（不改文件），返回 {部件: 次数}"""
    with zipfile.ZipFile(path) as z:
        names = z.namelist()
        targets = _target_parts(names, scope)
        out = {}
        for part in targets:
            xml = z.read(part).decode("utf-8")
            _, st = replace_in_xml(xml, old, case_sensitive, count_only=True)
            if st["found"]:
                out[part] = st["found"]
        return out
