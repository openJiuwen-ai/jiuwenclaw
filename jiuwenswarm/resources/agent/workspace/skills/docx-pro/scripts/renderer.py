# -*- coding: utf-8 -*-
"""
docx-pro 渲染引擎：主题系统 + 文档块渲染器。

被 docx_pro.py（create 子命令）与 md_parser.py（from-md 子命令）共用。
输入为统一的大纲（outline）字典：

{
  "title": "文档标题",
  "subtitle": "副标题（可选）",
  "meta": ["作者：张三", "日期：2026-01-01"],   # 封面/文档头元信息（可选）
  "cover": True,        # 是否生成独立封面页
  "toc": True,          # 是否在正文前插入目录域（Word 中按 F9 或右键更新域）
  "header": "页眉文字",  # 页眉（可选）
  "footer": "center",   # 页码样式：none | center | "X/Y"（可选，默认 none）
  "theme": "business",  # business | academic | minimal
  "watermark": "DRAFT", # 文字水印（可选）
  "sections": [ ...块列表... ]
}

支持的块（sections 内元素，均有 "type" 字段）：
  {"type":"heading","level":1..4,"text":"一、概述"}
  {"type":"para","text":"正文","indent":true,"align":"left|center|right|justify",
   "runs":[{"text":"..","bold":true,"italic":true,"code":true,"link":"url"},...]}
  {"type":"bullets","items":[".."|{"text":"..","bold_prefix":"..","indent":0|1}],"number":false}
  {"type":"table","header":[".."],"rows":[[".."|{"t":"..","cs":2,"rs":2}],...],
   "widths":[厘米...],"zebra":true,"align":"left|center"}
  {"type":"image","path":"a.png","width_cm":12.0,"caption":"图 1 xxx"}
  {"type":"code","text":"代码文本（可多行）"}
  {"type":"quote","text":"引用文本"}
  {"type":"divider"}
  {"type":"pagebreak"}
  {"type":"kv","items":{"键":"值",...}}   # 键值两列表格
"""
import os
import re
from xml.sax.saxutils import escape

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import parse_xml

# ---------------------------------------------------------------- 主题系统

THEMES = {
    "business": {  # 商务蓝：封面深蓝大标题，表头深蓝底白字，斑马纹浅蓝
        "title_color": "1F3B57", "subtitle_color": "2E74B5", "meta_color": "595959",
        "h1_color": "1F3B57", "h2_color": "2E74B5", "h3_color": "404040",
        "header_bg": "1F3B57", "zebra_bg": "EDF2F8", "quote_color": "595959",
        "code_bg": "F5F5F5", "h1_border": "1F3B57",
        "east_head": "微软雅黑", "east_body": "宋体", "latin": "Calibri", "mono": "Consolas",
    },
    "academic": {  # 学术黑白：衬线正文，表头黑底白字，无斑马纹
        "title_color": "000000", "subtitle_color": "404040", "meta_color": "595959",
        "h1_color": "000000", "h2_color": "000000", "h3_color": "333333",
        "header_bg": "000000", "zebra_bg": "F2F2F2", "quote_color": "595959",
        "code_bg": "F2F2F2", "h1_border": "000000",
        "east_head": "黑体", "east_body": "宋体", "latin": "Times New Roman", "mono": "Courier New",
    },
    "minimal": {  # 极简：无底色表头（加粗+粗底边框），大量留白
        "title_color": "111111", "subtitle_color": "666666", "meta_color": "888888",
        "h1_color": "111111", "h2_color": "333333", "h3_color": "444444",
        "header_bg": "FFFFFF", "zebra_bg": "FAFAFA", "quote_color": "777777",
        "code_bg": "F7F7F7", "h1_border": "111111",
        "east_head": "微软雅黑", "east_body": "微软雅黑", "latin": "Calibri", "mono": "Consolas",
    },
}

DEFAULT_THEME = "business"


def get_theme(name):
    return dict(THEMES.get(name, THEMES[DEFAULT_THEME]))


def _rgb(hexstr):
    return RGBColor(int(hexstr[0:2], 16), int(hexstr[2:4], 16), int(hexstr[4:6], 16))


# ---------------------------------------------------------------- 基础工具

def _set_font(run, theme, east=None, latin=None, size=None, bold=None,
              italic=None, color=None, mono=False):
    east = east or (theme["mono"] if mono else theme["east_body"])
    latin = latin or (theme["mono"] if mono else theme["latin"])
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = _rgb(color) if isinstance(color, str) else color


_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _shade(element_pr, fill):
    """给 pPr 或 tcPr 附加底纹。"""
    shd = element_pr.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    element_pr.append(shd)


def _para_bottom_border(p, color, sz="8"):
    ppr = p._element.get_or_add_pPr()
    pbdr = ppr.makeelement(qn("w:pBdr"), {})
    bottom = ppr.makeelement(qn("w:bottom"), {})
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    ppr.append(pbdr)


INLINE_TOKEN_RE = re.compile(
    r"(\*\*(?:[^*]|\*(?!\*))+?\*\*|`[^`]+?`|~~[^~]+?~~|"
    r"\[[^\]]+?\]\([^)]+?\)|\*[^*\s][^*]*?\*)"
)


def parse_inline(text):
    """把 **粗体** / `代码` / ~~删除线~~ / [文字](链接) / *斜体* 解析为 run 列表。"""
    runs = []
    pos = 0
    for m in INLINE_TOKEN_RE.finditer(text):
        if m.start() > pos:
            runs.append({"text": text[pos:m.start()]})
        tok = m.group(0)
        if tok.startswith("**"):
            runs.append({"text": tok[2:-2], "bold": True})
        elif tok.startswith("`"):
            runs.append({"text": tok[1:-1], "code": True})
        elif tok.startswith("~~"):
            runs.append({"text": tok[2:-2], "strike": True})
        elif tok.startswith("["):
            inner = re.match(r"\[([^\]]+)\]\(([^)]+)\)", tok)
            runs.append({"text": inner.group(1), "link": inner.group(2)})
        else:
            runs.append({"text": tok[1:-1], "italic": True})
        pos = m.end()
    if pos < len(text):
        runs.append({"text": text[pos:]})
    return runs or [{"text": text}]


def add_runs(p, spec, theme, size=11, color=None, bold=None):
    """向段落写入行内 run 列表（支持 bold/italic/code/strike/link）。"""
    runs = spec if isinstance(spec, list) else parse_inline(spec)
    for r in runs:
        text = r.get("text", "")
        if not text:
            continue
        run = p.add_run(text)
        _set_font(run, theme, size=size,
                  bold=r.get("bold", bold),
                  italic=r.get("italic"),
                  color=r.get("code") and "C7254E" or color,
                  mono=r.get("code", False))
        if r.get("strike"):
            run.font.strike = True
        if r.get("link"):
            add_hyperlink(p, r["link"], text, theme, size)


def add_hyperlink(p, url, text, theme, size=11):
    """在段落中插入真正的超链接（非纯文本）。"""
    part = p.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hyperlink = parse_xml(
        '<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'r:id="%s" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        "</w:hyperlink>" % r_id)
    new_run = parse_xml(
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:rPr><w:color w:val="0563C1"/><w:u w:val="single"/></w:rPr>'
        "<w:t>%s</w:t></w:r>" % escape(text))
    hyperlink.append(new_run)
    p._p.append(hyperlink)


# ---------------------------------------------------------------- 域（field）

def _field_run(p, instr, theme, size=10):
    run = p.add_run()
    _set_font(run, theme, size=size)
    r = run._element
    fld1 = r.makeelement(qn("w:fldChar"), {})
    fld1.set(qn("w:fldCharType"), "begin")
    instr_el = r.makeelement(qn("w:instrText"), {})
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    fld2 = r.makeelement(qn("w:fldChar"), {})
    fld2.set(qn("w:fldCharType"), "end")
    r.append(fld1)
    r.append(instr_el)
    r.append(fld2)
    return run


def add_toc(doc, theme, levels="1-3"):
    """插入目录域。Word 打开后按 F9（或右键→更新域）生成实际目录。"""
    p = doc.add_paragraph()
    run = p.add_run()
    _set_font(run, theme, size=10.5, color="808080")
    r = run._element
    fld1 = r.makeelement(qn("w:fldChar"), {})
    fld1.set(qn("w:fldCharType"), "begin")
    instr = r.makeelement(qn("w:instrText"), {})
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "%s" \\h \\z \\u' % levels
    fld_sep = r.makeelement(qn("w:fldChar"), {})
    fld_sep.set(qn("w:fldCharType"), "separate")
    t = r.makeelement(qn("w:t"), {})
    t.text = "【目录域已插入：在 Word 中选中此处按 F9 或右键“更新域”即可生成目录】"
    fld_end = r.makeelement(qn("w:fldChar"), {})
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld1, instr, fld_sep, t, fld_end):
        r.append(el)
    return p


def add_page_number_footer(doc, theme, style="center"):
    """页脚页码。style: 'center' -> 第 X 页；'X/Y' -> X / Y。"""
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if style == "center":
        pre = p.add_run("第 ")
        _set_font(pre, theme, size=9, color="808080")
        _field_run(p, "PAGE", theme)
        post = p.add_run(" 页")
        _set_font(post, theme, size=9, color="808080")
    else:
        _field_run(p, "PAGE", theme)
        sep = p.add_run(" / ")
        _set_font(sep, theme, size=9, color="808080")
        _field_run(p, "NUMPAGES", theme)
    for run in p.runs:
        if run.font.size is None:
            _set_font(run, theme, size=9, color="808080")


def add_header(doc, theme, text, align="center"):
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = _ALIGN.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text)
    _set_font(run, theme, east=theme["east_head"], size=9, color="808080")
    _para_bottom_border(p, "CCCCCC", sz="4")


# ---------------------------------------------------------------- 水印

def add_watermark(doc, text, color="C0C0C0", opacity=0.45):
    """整页文字水印（利用页眉 VML shape 实现，所有页面生效）。"""
    for i, section in enumerate(doc.sections):
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        xml = (
            '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<w:pict>"
            '<v:shape xmlns:v="urn:schemas-microsoft-com:vml" '
            'xmlns:o="urn:schemas-microsoft-com:office:office" '
            'id="WordPictureWatermark%d" o:spid="_x0000_s20%02d" type="#_x0000_t136" '
            'style="position:absolute;margin-left:0;margin-top:0;width:420pt;height:180pt;'
            "rotation:315;z-index:-251658752;mso-position-horizontal:center;"
            "mso-position-horizontal-relative:margin;mso-position-vertical:center;"
            'mso-position-vertical-relative:margin" o:allowincell="f" '
            'fillcolor="#%s" stroked="f">'
            '<v:fill opacity="%.2f"/>'
            '<v:textpath style="font-family:&quot;\u5fae\u8f6f\u96c5\u9ed1&quot;;'
            'font-size:1pt" string="%s"/>'
            "</v:shape></w:pict></w:r>"
            % (i, 49 + i, color, opacity, escape(text))
        )
        p._p.append(parse_xml(xml))


# ---------------------------------------------------------------- 块渲染器

def render_heading(doc, block, theme):
    level = int(block.get("level", 1))
    level = max(1, min(4, level))
    sizes = {1: 15, 2: 13, 3: 11.5, 4: 11}
    colors = {1: theme["h1_color"], 2: theme["h2_color"], 3: theme["h3_color"], 4: theme["h3_color"]}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt({1: 16, 2: 10, 3: 8, 4: 6}[level])
    p.paragraph_format.space_after = Pt({1: 8, 2: 5, 3: 4, 4: 3}[level])
    p.paragraph_format.keep_with_next = True
    add_runs(p, block.get("text", ""), theme,
             size=sizes[level], bold=True, color=colors[level])
    for run in p.runs:  # 标题整体用标题字体
        _set_font(run, theme, east=theme["east_head"],
                  latin=theme["latin"], size=sizes[level],
                  bold=True, color=colors[level])
    if level == 1:
        _para_bottom_border(p, theme["h1_border"], sz="8")
    # 映射到 Word 内置 Heading 样式（供 TOC / 导航窗格识别）
    p.style = doc.styles["Heading %d" % level]
    for run in p.runs:
        _set_font(run, theme, east=theme["east_head"],
                  latin=theme["latin"], size=sizes[level],
                  bold=True, color=colors[level])
    return p


def render_para(doc, block, theme):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.4
    pf.space_after = Pt(4)
    if block.get("indent", True):
        pf.first_line_indent = Pt(21)
    if block.get("align"):
        p.alignment = _ALIGN.get(block["align"], WD_ALIGN_PARAGRAPH.LEFT)
    if "runs" in block:
        add_runs(p, block["runs"], theme, size=11)
    else:
        add_runs(p, block.get("text", ""), theme, size=11)
    return p


def render_bullets(doc, block, theme):
    items = block.get("items", [])
    out = []
    for item in items:
        if isinstance(item, dict):
            text = item.get("text", "")
            prefix = item.get("bold_prefix")
            indent = int(item.get("indent", 0))
        else:
            text, prefix, indent = item, None, 0
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.35
        pf.space_after = Pt(3)
        pf.left_indent = Pt(21 + 21 * indent)
        marker = p.add_run("%s " % (block.get("number") and "1." or "\u2022"))
        _set_font(marker, theme, size=11, color=theme["h2_color"], bold=True)
        if prefix:
            prun = p.add_run(prefix)
            _set_font(prun, theme, size=11, bold=True, color=theme["h1_color"])
        add_runs(p, text, theme, size=11)
        out.append(p)
    return out


def _grid_positions(rows_spec, ncols):
    """计算每个 cell 的 (行, 起始列)，处理 rs/cs 占位。"""
    nrows = len(rows_spec)
    occupied = [[False] * ncols for _ in range(nrows)]
    positions = []
    for r, row in enumerate(rows_spec):
        c = 0
        for cell in row:
            while c < ncols and occupied[r][c]:
                c += 1
            if c >= ncols:
                break
            cs = int(cell.get("cs", 1)) if isinstance(cell, dict) else 1
            rs = int(cell.get("rs", 1)) if isinstance(cell, dict) else 1
            for dr in range(rs):
                for dc in range(cs):
                    rr, cc = r + dr, c + dc
                    if rr < nrows and cc < ncols:
                        occupied[rr][cc] = True
            positions.append((r, c, cell))
            c += cs
    return positions


def _cell_text(cell):
    if isinstance(cell, dict):
        return cell.get("t", cell.get("text", ""))
    return cell


def render_table(doc, block, theme):
    header = block.get("header") or []
    rows = block.get("rows") or []
    ncols = max([len(header)] + [len(r) for r in rows]) if (header or rows) else 1
    ncols = int(block.get("ncols", ncols))
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    minimal = theme is THEMES.get("minimal") or theme.get("header_bg") == "FFFFFF"

    def style_cell(cell, text, is_header, row_idx, align):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = _ALIGN.get(align, WD_ALIGN_PARAGRAPH.LEFT) if not is_header \
            else WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        run = p.add_run(str(text))
        if is_header:
            _set_font(run, theme, east=theme["east_head"], size=10.5, bold=True,
                      color="FFFFFF" if not minimal else "111111")
        else:
            _set_font(run, theme, size=10.5)

    # 表头
    if header:
        hdr_cells = table.add_row().cells
        grid = _grid_positions([header], ncols)
        merged = set()
        for (r, c, cell) in grid:
            spec = cell if isinstance(cell, dict) else {"t": cell}
            cs, rs = int(spec.get("cs", 1)), int(spec.get("rs", 1))
            if cs > 1 or rs > 1:
                hdr_cells[c].merge(hdr_cells[min(c + cs - 1, ncols - 1)])
                merged.add(c)
        for (r, c, cell) in grid:
            style_cell(table.rows[0].cells[c], _cell_text(cell), True, 0,
                       block.get("align", "center"))
        for c in range(ncols):  # 表头底纹/边框
            tc = table.rows[0].cells[c]
            tcpr = tc._element.get_or_add_tcPr()
            if not minimal:
                _shade(tcpr, theme["header_bg"])
            else:
                tcpr.append(parse_xml(
                    '<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/'
                    'wordprocessingml/2006/main">'
                    '<w:bottom w:val="single" w:sz="12" w:color="%s"/></w:tcBorders>'
                    % theme["h1_border"]))

    # 数据行（注意：表头已占 1 行，数据行需偏移 header_rows）
    header_rows = 1 if header else 0
    zebra = block.get("zebra", True)
    align = block.get("align", "left")
    if rows:
        grid = _grid_positions(rows, ncols)
        nrows = max(r + (int(cell.get("rs", 1)) if isinstance(cell, dict) else 1)
                    for r, c, cell in grid) if grid else len(rows)
        for _ in range(nrows):
            table.add_row()
        for (r, c, cell) in grid:
            spec = cell if isinstance(cell, dict) else {"t": cell}
            cs, rs = int(spec.get("cs", 1)), int(spec.get("rs", 1))
            if cs > 1 or rs > 1:
                r2 = min(r + rs - 1, nrows - 1) + header_rows
                c2 = min(c + cs - 1, ncols - 1)
                table.cell(r + header_rows, c).merge(table.cell(r2, c2))
        for (r, c, cell) in grid:
            style_cell(table.cell(r + header_rows, c), _cell_text(cell),
                       False, r, align)
        if zebra:
            for tr in range(header_rows, len(table.rows)):
                if (tr - header_rows) % 2 == 1:  # 数据行第 2、4… 行着色
                    for c in range(ncols):
                        _shade(table.rows[tr].cells[c]._element.get_or_add_tcPr(),
                               theme["zebra_bg"])

    widths = block.get("widths")
    if widths:
        for i, w in enumerate(widths):
            if i < ncols:
                for row in table.rows:
                    row.cells[i].width = Cm(float(w))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def render_kv(doc, block, theme):
    items = block.get("items") or {}
    rows = [[k, v] for k, v in items.items()]
    return render_table(doc, {"header": None, "rows": rows, "zebra": True,
                              "widths": [4.5, 11.5]}, theme)


def render_image(doc, block, theme):
    path = block.get("path", "")
    if not os.path.exists(path):
        raise FileNotFoundError("图片不存在: %s" % path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    width = block.get("width_cm")
    run.add_picture(path, width=Cm(float(width)) if width else None)
    caption = block.get("caption")
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(8)
        crun = cap.add_run(caption)
        _set_font(crun, theme, east=theme["east_head"], size=9, color="808080")
    return p


def render_code(doc, block, theme):
    text = block.get("text", "")
    lines = text.split("\n")
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(6 if i == 0 else 0)
        pf.space_after = Pt(6 if i == len(lines) - 1 else 0)
        pf.line_spacing = 1.15
        pf.left_indent = Pt(14)
        pf.right_indent = Pt(14)
        run = p.add_run(line if line else " ")
        _set_font(run, theme, east=theme["mono"], size=9.5, color="333333", mono=True)
        _shade(p._element.get_or_add_pPr(), theme["code_bg"])
    return None


def render_quote(doc, block, theme):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Pt(14)
    pf.right_indent = Pt(14)
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.35
    add_runs(p, block.get("text", ""), theme, size=10.5, color=theme["quote_color"])
    ppr = p._element.get_or_add_pPr()
    pbdr = ppr.makeelement(qn("w:pBdr"), {})
    left = ppr.makeelement(qn("w:left"), {})
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), theme["h2_color"])
    pbdr.append(left)
    ppr.append(pbdr)
    return p


def render_divider(doc, block, theme):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    _para_bottom_border(p, "CCCCCC", sz="6")
    return p


def render_pagebreak(doc, block, theme):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    return p


BLOCK_RENDERERS = {
    "heading": render_heading,
    "para": render_para,
    "p": render_para,
    "bullets": render_bullets,
    "list": render_bullets,
    "table": render_table,
    "kv": render_kv,
    "image": render_image,
    "img": render_image,
    "code": render_code,
    "quote": render_quote,
    "divider": render_divider,
    "hr": render_divider,
    "pagebreak": render_pagebreak,
    "page_break": render_pagebreak,
}


# ---------------------------------------------------------------- 文档级渲染

def _setup_page(doc):
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(2.3)
    section.left_margin = section.right_margin = Cm(2.6)


def _setup_normal_style(doc, theme):
    normal = doc.styles["Normal"]
    normal.font.name = theme["latin"]
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), theme["east_body"])


def _render_cover(doc, outline, theme):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(150)
    run = p.add_run(outline.get("title", ""))
    _set_font(run, theme, east=theme["east_head"], size=30, bold=True,
              color=theme["title_color"])
    if outline.get("subtitle"):
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.paragraph_format.space_before = Pt(12)
        srun = sp.add_run(outline["subtitle"])
        _set_font(srun, theme, east=theme["east_head"], size=15,
                  color=theme["subtitle_color"])
    for m in outline.get("meta", []):
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mp.paragraph_format.space_before = Pt(6)
        mrun = mp.add_run(m)
        _set_font(mrun, theme, size=11, color=theme["meta_color"])
    render_pagebreak(doc, None, theme)


def _render_doc_header(doc, outline, theme):
    """无封面时在文档开头渲染大标题块。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(outline.get("title", ""))
    _set_font(run, theme, east=theme["east_head"], size=24, bold=True,
              color=theme["title_color"])
    if outline.get("subtitle"):
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.paragraph_format.space_after = Pt(4)
        srun = sp.add_run(outline["subtitle"])
        _set_font(srun, theme, east=theme["east_head"], size=13,
                  color=theme["subtitle_color"])
    if outline.get("meta"):
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mp.paragraph_format.space_after = Pt(14)
        mrun = mp.add_run("    ".join(outline["meta"]))
        _set_font(mrun, theme, size=10, color=theme["meta_color"])
    _para_bottom_border(p, theme["h1_border"], sz="8")


def render_outline(outline, out_path):
    """把大纲字典渲染为 docx 文件。"""
    theme = get_theme(outline.get("theme", DEFAULT_THEME))
    doc = Document()
    _setup_page(doc)
    _setup_normal_style(doc, theme)

    if outline.get("cover"):
        _render_cover(doc, outline, theme)
    elif outline.get("title"):
        _render_doc_header(doc, outline, theme)

    if outline.get("toc"):
        tp = doc.add_paragraph()
        trun = tp.add_run("目  录")
        _set_font(trun, theme, east=theme["east_head"], size=16, bold=True,
                  color=theme["h1_color"])
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp.paragraph_format.space_before = Pt(6)
        tp.paragraph_format.space_after = Pt(10)
        add_toc(doc, theme, outline.get("toc_levels", "1-3"))
        if outline.get("cover"):  # 有封面时目录单独成页更美观
            render_pagebreak(doc, None, theme)

    if outline.get("header"):
        add_header(doc, theme, outline["header"])
    footer_style = outline.get("footer")
    if footer_style and footer_style != "none":
        add_page_number_footer(doc, theme,
                               "center" if footer_style is True else footer_style)

    for block in outline.get("sections", []):
        btype = block.get("type", "para")
        renderer = BLOCK_RENDERERS.get(btype)
        if renderer is None:
            raise ValueError("未知的块类型: %r（支持: %s）"
                             % (btype, ", ".join(sorted(BLOCK_RENDERERS))))
        renderer(doc, block, theme)

    if outline.get("watermark"):
        add_watermark(doc, str(outline["watermark"]))

    doc.save(out_path)
    return out_path
