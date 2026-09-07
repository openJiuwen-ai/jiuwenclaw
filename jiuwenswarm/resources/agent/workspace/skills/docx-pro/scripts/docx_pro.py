# -*- coding: utf-8 -*-
"""
docx_pro.py - 专业 Word 文档操作命令行工具（docx-pro 技能主入口）

子命令：
  create     从 JSON 大纲生成富格式 docx（封面/目录/多级标题/表格/图片/代码块/引用/页眉页脚/水印）
  from-md    Markdown 转 docx（复用 create 的渲染引擎与主题）
  to-md      docx 转 Markdown（可选导出图片）
  replace    docx 保真文本替换（跨 run 定位、表格/页眉页脚覆盖、其余字节零改动、批量映射）
  inspect    查看文档结构（标题树/表格/图片/页面信息，支持 --json）
  toc        为已有 docx 插入目录域（Word 中按 F9 更新生成目录）
  watermark  为已有 docx 添加文字水印

用法示例：
  python docx_pro.py create -j outline.json -o 报告.docx --theme business
  python docx_pro.py from-md 输入.md -o 输出.docx --theme academic --toc
  python docx_pro.py to-md 输入.docx -o 输出.md --images
  python docx_pro.py replace 合同.docx -o 合同_新.docx -f "甲方：A 公司" -t "甲方：B公司"
  python docx_pro.py replace 合同.docx -o 合同_新.docx --map pairs.json --dry-run
  python docx_pro.py inspect 输入.docx
  python docx_pro.py toc 输入.docx -o 带目录.docx
  python docx_pro.py watermark 输入.docx -o 加水印.docx -t "DRAFT"
"""
import io
import json
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document                    # noqa: E402
from docx.shared import Pt, Cm               # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml.ns import qn                  # noqa: E402

import renderer                               # noqa: E402
import docx_replace                           # noqa: E402
from md_parser import parse_markdown          # noqa: E402
from md_export import docx_to_markdown        # noqa: E402


def _ensure_utf8_console():
    """控制台编码兜底：保持终端原编码（如 GBK，中文可正常显示），
    但启用 errors=replace，避免个别字符触发 UnicodeEncodeError。"""
    for name in ("stdout", "stderr"):
        stream = getattr(sys, name, None)
        if stream is None or not hasattr(stream, "buffer"):
            continue
        try:
            enc = stream.encoding or "utf-8"
            setattr(sys, name, io.TextIOWrapper(
                stream.buffer, encoding=enc, errors="replace"))
        except Exception:
            pass


# ---------------------------------------------------------------- create

def cmd_create(args):
    if args.json == "-":
        outline = json.loads(sys.stdin.read())
    else:
        with open(args.json, encoding="utf-8") as f:
            outline = json.loads(f.read())
    for key in ("theme", "header", "footer", "toc", "cover", "watermark"):
        val = getattr(args, key, None)
        if val is None:
            continue
        if key in ("toc", "cover") and isinstance(val, str):
            val = val.strip().lower() in ("true", "yes", "1", "")
        outline[key] = val
    if not outline.get("title"):
        outline["title"] = os.path.splitext(os.path.basename(args.out))[0]
    if os.path.abspath(args.out) == os.path.abspath(args.json):
        print("[错误] 输出文件不得覆盖输入的 JSON 文件")
        return 2
    # 图片相对路径：相对于大纲 JSON 所在目录解析
    base_dir = os.path.dirname(os.path.abspath(args.json)) if args.json != "-" \
        else os.getcwd()
    for block in outline.get("sections", []):
        if block.get("type") in ("image", "img") and block.get("path") \
                and not os.path.isabs(block["path"]):
            block["path"] = os.path.join(base_dir, block["path"])
    renderer.render_outline(outline, args.out)
    print("[OK] 已生成: %s（主题: %s，块数: %d）"
          % (args.out, outline.get("theme", renderer.DEFAULT_THEME),
             len(outline.get("sections", []))))
    return 0


# ---------------------------------------------------------------- from-md

def cmd_from_md(args):
    with open(args.md, encoding="utf-8") as f:
        text = f.read()
    outline = parse_markdown(text, base_dir=os.path.dirname(os.path.abspath(args.md)))
    if args.theme:
        outline["theme"] = args.theme
    if args.toc:
        outline["toc"] = True
    if args.cover:
        outline["cover"] = True
    if args.header:
        outline["header"] = args.header
    if args.footer:
        outline["footer"] = args.footer
    if args.watermark:
        outline["watermark"] = args.watermark
    if not outline.get("title"):
        outline["title"] = os.path.splitext(os.path.basename(args.out))[0]
    renderer.render_outline(outline, args.out)
    print("[OK] Markdown 已转换为: %s（主题: %s，块数: %d）"
          % (args.out, outline.get("theme"), len(outline.get("sections", []))))
    return 0


# ---------------------------------------------------------------- to-md

def cmd_to_md(args):
    images_dir = None
    if args.images:
        images_dir = os.path.join(
            os.path.dirname(os.path.abspath(args.out)) or ".",
            os.path.splitext(os.path.basename(args.out))[0] + "_images")
    md = docx_to_markdown(args.docx, images_dir)
    with open(args.out, "w", encoding="utf-8") as f:
        f.write(md)
    extra = ""
    if images_dir:
        extra = "，图片目录: %s" % images_dir
    print("[OK] docx 已转换为: %s%s" % (args.out, extra))
    return 0


# ---------------------------------------------------------------- replace

def _load_replace_pairs(args):
    """整合 -f/-t 与 --map 两种来源的替换对"""
    pairs = []
    if args.map:
        with open(args.map, encoding="utf-8") as f:
            data = json.load(f)
        if isinstance(data, dict):
            pairs = [[str(k), str(v)] for k, v in data.items()]
        elif isinstance(data, list):
            for it in data:
                if isinstance(it, (list, tuple)) and len(it) == 2:
                    pairs.append([str(it[0]), str(it[1])])
                elif isinstance(it, dict) and "find" in it and "to" in it:
                    pairs.append([str(it["find"]), str(it["to"])])
                else:
                    raise ValueError("--map 条目格式应为 [旧, 新] 或 {find, to}")
        else:
            raise ValueError("--map 仅接受 JSON 数组或对象")
    if args.find or args.to:
        finds = args.find or []
        tos = args.to or []
        if len(finds) != len(tos):
            raise ValueError("--find 与 --to 必须成对出现（当前 %d 对 %d）"
                             % (len(finds), len(tos)))
        pairs = [[f, t] for f, t in zip(finds, tos)] + pairs
    if not pairs:
        raise ValueError("请通过 -f/-t 或 --map 提供至少一组替换")
    for old, new in pairs:
        if old == new:
            raise ValueError("查找与替换文本相同：%r" % old)
    return pairs


def cmd_replace(args):
    pairs = _load_replace_pairs(args)
    if os.path.abspath(args.out) == os.path.abspath(args.docx):
        print("[错误] 输出文件不得覆盖输入文件")
        return 2

    report = docx_replace.replace_in_docx(
        args.docx, args.out, pairs, scope=args.scope, dry_run=args.dry_run)

    verb = "预览" if args.dry_run else "替换完成"
    total = report["total_found"]
    print("[%s] %s：共 %d 处命中%s" % (
        "OK" if total else "警告", verb, total,
        "" if not args.dry_run else "（未写出文件）"))
    for old, new, per_part in report["pairs"]:
        parts_desc = "，".join(
            "%s: %d" % (p, st["found"]) for p, st in sorted(per_part.items())
            if st["found"])
        cross = sum(st["cross_run"] for st in per_part.values())
        cross_note = "（含跨 run %d 处）" % cross if cross else ""
        print("  「%s」→「%s」：%d 处%s%s"
              % (old, new, sum(st["found"] for st in per_part.values()),
                 cross_note, ("  [%s]" % parts_desc) if parts_desc else ""))
    if args.dry_run:
        return 0 if total else 1

    v = report
    changed = "、".join(v.get("changed_entries", []))
    n_entries = len(v.get("parts_scanned", []))
    print("  已写出: %s" % args.out)
    print("  变化部件: %s；其余 zip 条目字节级一致: %s"
          % (changed or "（无）", v.get("other_entries_byte_identical")))
    if v.get("opens_ok") is True:
        print("  复检: 段落 %d/%d，表格 %d/%d，新文件可正常打开"
              % (v["paragraphs"][0], v["paragraphs"][1],
                 v["tables"][0], v["tables"][1]))
    residual = {k: c for k, c in v.get("residual", {}).items() if c}
    if residual:
        print("  [提示] 残留 %s（若新文本本身包含旧文本属正常）" % residual)
    if not v.get("entry_list_identical", True) or v.get("unexpected_changes"):
        print("[警告] 校验异常：%s" % v.get("unexpected_changes"))
        return 1
    return 0 if total else 1


# ---------------------------------------------------------------- inspect

def _doc_info(path):
    doc = Document(path)
    headings = []
    paras = 0
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        paras += 1
        style = p.style.name if p.style is not None else ""
        if style and style.startswith("Heading"):
            m = re.match(r"Heading (\d)", style)
            if m:
                headings.append({"level": int(m.group(1)),
                                 "text": p.text.strip()[:60]})
    table_shapes = []
    for t in doc.tables:
        table_shapes.append({"rows": len(t.rows), "cols": len(t.columns)})
    section = doc.sections[0]
    try:
        page = "%.0f x %.0f mm" % (section.page_width.mm, section.page_height.mm)
        margins = "上%.0f/下%.0f/左%.0f/右%.0f mm" % (
            section.top_margin.mm, section.bottom_margin.mm,
            section.left_margin.mm, section.right_margin.mm)
    except Exception:
        page = margins = "未知"
    return {
        "file": path,
        "paragraphs_nonempty": paras,
        "headings": headings,
        "tables": {"count": len(table_shapes), "shapes": table_shapes},
        "images": len(doc.inline_shapes),
        "page": page,
        "margins": margins,
    }


def cmd_inspect(args):
    info = _doc_info(args.docx)
    if args.json:
        print(json.dumps(info, ensure_ascii=False, indent=2))
        return 0
    print("文件: %s" % info["file"])
    print("页面: %s（边距 %s）" % (info["page"], info["margins"]))
    print("非空段落数: %d | 表格: %d | 内联图片: %d"
          % (info["paragraphs_nonempty"], info["tables"]["count"], info["images"]))
    if info["headings"]:
        print("标题树:")
        for h in info["headings"]:
            print("  %s%s  %s" % ("    " * (h["level"] - 1), "#" * h["level"], h["text"]))
    for i, s in enumerate(info["tables"]["shapes"][:10], 1):
        print("表格%d: %d 行 x %d 列" % (i, s["rows"], s["cols"]))
    return 0


# ---------------------------------------------------------------- toc

def cmd_toc(args):
    doc = Document(args.docx)
    theme = renderer.get_theme(args.theme)
    # 构造目录标题段与目录域段
    toc_title = doc.add_paragraph()
    trun = toc_title.add_run("目  录")
    renderer._set_font(trun, theme, east=theme["east_head"], size=16, bold=True,
                       color=theme["h1_color"])
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_title.paragraph_format.space_before = Pt(6)
    toc_title.paragraph_format.space_after = Pt(10)

    toc_p = doc.add_paragraph()
    run = toc_p.add_run()
    r = run._element
    fld1 = r.makeelement(qn("w:fldChar"), {})
    fld1.set(qn("w:fldCharType"), "begin")
    instr = r.makeelement(qn("w:instrText"), {})
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "%s" \\h \\z \\u' % args.levels
    sep = r.makeelement(qn("w:fldChar"), {})
    sep.set(qn("w:fldCharType"), "separate")
    t = r.makeelement(qn("w:t"), {})
    t.text = "【目录域已插入：在 Word 中选中此处按 F9 或右键\u201c更新域\u201d即可生成目录】"
    end = r.makeelement(qn("w:fldChar"), {})
    end.set(qn("w:fldCharType"), "end")
    for el in (fld1, instr, sep, t, end):
        r.append(el)

    # 移动到正文最前（sectPr 之前）
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    anchor = sect_pr if sect_pr is not None else list(body)[-1]
    for el in (toc_title._p, toc_p._p):
        body.remove(el)
        anchor.addprevious(el)
    doc.save(args.out)
    print("[OK] 已插入目录域: %s（标题层级: %s，Word 中按 F9 更新）"
          % (args.out, args.levels))
    return 0


# ---------------------------------------------------------------- watermark

def cmd_watermark(args):
    doc = Document(args.docx)
    renderer.add_watermark(doc, args.text, color=args.color, opacity=args.opacity)
    doc.save(args.out)
    print("[OK] 已添加水印 \u201c%s\u201d: %s" % (args.text, args.out))
    return 0


# ---------------------------------------------------------------- main

def build_parser():
    import argparse
    ap = argparse.ArgumentParser(
        prog="docx_pro.py",
        description="docx-pro: 专业 Word 文档操作工具（生成/转换/保真替换/分析/目录/水印）")
    sub = ap.add_subparsers(dest="cmd", required=True)

    c = sub.add_parser("create", help="从 JSON 大纲生成富格式 docx")
    c.add_argument("-j", "--json", required=True, help="大纲 JSON 文件路径（- 表示 stdin）")
    c.add_argument("-o", "--out", required=True, help="输出 docx 路径（不得覆盖输入）")
    c.add_argument("--theme", choices=list(renderer.THEMES),
                   help="主题：business/academic/minimal")
    for key, help_ in (("header", "页眉文字"), ("footer", "页码样式 center 或 X/Y"),
                       ("toc", "插入目录域(true)"), ("cover", "生成封面页(true)"),
                       ("watermark", "文字水印")):
        c.add_argument("--" + key, help=help_)
    c.set_defaults(func=cmd_create)

    fm = sub.add_parser("from-md", help="Markdown 转 docx")
    fm.add_argument("md", help="输入 Markdown 文件")
    fm.add_argument("-o", "--out", required=True, help="输出 docx 路径")
    fm.add_argument("--theme", choices=list(renderer.THEMES))
    fm.add_argument("--toc", action="store_true", help="插入目录域")
    fm.add_argument("--cover", action="store_true", help="生成封面页")
    fm.add_argument("--header", help="页眉文字")
    fm.add_argument("--footer", help="页码样式 center 或 X/Y")
    fm.add_argument("--watermark", help="文字水印")
    fm.set_defaults(func=cmd_from_md)

    tm = sub.add_parser("to-md", help="docx 转 Markdown")
    tm.add_argument("docx", help="输入 docx 文件")
    tm.add_argument("-o", "--out", required=True, help="输出 Markdown 路径")
    tm.add_argument("--images", action="store_true", help="导出文档内图片")
    tm.set_defaults(func=cmd_to_md)

    rp = sub.add_parser(
        "replace", help="docx 保真文本替换（跨 run 定位，格式零改动）")
    rp.add_argument("docx", help="输入 docx 文件")
    rp.add_argument("-o", "--out", required=True, help="输出 docx 路径（不得覆盖输入）")
    rp.add_argument("-f", "--find", action="append",
                    help="查找文本（与 --to 成对，可多次）")
    rp.add_argument("-t", "--to", action="append",
                    help="替换文本（与 --find 成对，可多次）")
    rp.add_argument("--map", help="批量替换映射 JSON：[[旧,新],...] 或 {旧:新}")
    rp.add_argument("--scope", choices=["all", "body"], default="all",
                    help="替换范围：all=正文+页眉页脚+脚注尾注批注（默认），body=仅正文")
    rp.add_argument("--dry-run", action="store_true",
                    help="仅统计出现次数，不写出文件")
    rp.set_defaults(func=cmd_replace)

    ins = sub.add_parser("inspect", help="查看文档结构")
    ins.add_argument("docx", help="输入 docx 文件")
    ins.add_argument("--json", action="store_true", help="以 JSON 输出")
    ins.set_defaults(func=cmd_inspect)

    tc = sub.add_parser("toc", help="为已有 docx 插入目录域")
    tc.add_argument("docx", help="输入 docx 文件")
    tc.add_argument("-o", "--out", required=True, help="输出 docx 路径")
    tc.add_argument("--levels", default="1-3", help="目录标题层级，如 1-3")
    tc.add_argument("--theme", default="business",
                    help="目录标题配色主题（仅影响目录标题字体颜色）")
    tc.set_defaults(func=cmd_toc)

    wm = sub.add_parser("watermark", help="为已有 docx 添加文字水印")
    wm.add_argument("docx", help="输入 docx 文件")
    wm.add_argument("-o", "--out", required=True, help="输出 docx 路径")
    wm.add_argument("-t", "--text", required=True, help="水印文字")
    wm.add_argument("--color", default="C0C0C0", help="水印颜色（6 位 HEX，默认 C0C0C0）")
    wm.add_argument("--opacity", type=float, default=0.45, help="不透明度 0-1，默认 0.45")
    wm.set_defaults(func=cmd_watermark)
    return ap


def main(argv=None):
    _ensure_utf8_console()
    ap = build_parser()
    args = ap.parse_args(argv)
    for path_attr in ("json", "md", "docx", "map"):
        path = getattr(args, path_attr, None)
        if path and path != "-" and not os.path.exists(path):
            print("[错误] 文件不存在: %s" % path)
            return 2
    try:
        return args.func(args)
    except FileNotFoundError as e:
        print("[错误] %s" % e)
        return 2
    except ValueError as e:
        print("[错误] %s" % e)
        return 2


if __name__ == "__main__":
    sys.exit(main())
