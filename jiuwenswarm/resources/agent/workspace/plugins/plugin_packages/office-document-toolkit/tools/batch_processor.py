import os
from pathlib import Path

from openjiuwen.core.foundation.tool import Tool, ToolCard

from pdf_font_utils import select_pdf_font


class BatchProcessor(Tool):
    """批量处理工具：批量重命名、批量合并、批量转格式、遍历文件夹。"""

    def __init__(self) -> None:
        super().__init__(
            ToolCard(
                id="batch_processor",
                name="batch_processor",
                description=(
                    "批量处理工具：批量重命名、批量合并多个Excel/Word、"
                    "批量转格式、遍历文件夹处理。"
                    "当用户需要批量处理多个文件时调用。"
                ),
                input_params={
                    "type": "object",
                    "properties": {
                        "operation": {
                            "type": "string",
                            "enum": [
                                "rename",
                                "merge_excel",
                                "merge_word",
                                "convert",
                                "list_files",
                            ],
                            "description": "批量操作类型",
                        },
                        "input_dir": {
                            "type": "string",
                            "description": "输入目录路径",
                        },
                        "pattern": {
                            "type": "string",
                            "description": "文件匹配模式，如 *.xlsx",
                        },
                        "options": {
                            "type": "object",
                            "description": (
                                "操作特定选项：rename(prefix/start_num/padding)、"
                                "convert(target_format)"
                            ),
                        },
                        "output_subdir": {
                            "type": "string",
                            "description": "输出子目录名，默认为 batch_output",
                        },
                    },
                    "required": ["operation", "input_dir"],
                },
            )
        )

    async def invoke(self, inputs, **kwargs):
        operation = inputs.get("operation", "")
        input_dir = inputs.get("input_dir", "")
        pattern = inputs.get("pattern", "*")
        options = inputs.get("options", {})
        output_subdir = inputs.get("output_subdir", "batch_output")

        if not operation:
            return {"success": False, "error": "缺少 operation 参数"}
        if not input_dir or not os.path.isdir(input_dir):
            return {
                "success": False,
                "error": f"输入目录不存在: {input_dir}",
            }

        from openjiuwen.core.sys_operation.cwd import get_cwd

        base_dir = Path(get_cwd()) / output_subdir
        base_dir.mkdir(parents=True, exist_ok=True)

        try:
            if operation == "rename":
                result = self._batch_rename(input_dir, pattern, options)
            elif operation == "merge_excel":
                result = self._merge_excel_files(
                    input_dir, pattern, str(base_dir)
                )
            elif operation == "merge_word":
                result = self._merge_word_files(
                    input_dir, pattern, str(base_dir)
                )
            elif operation == "convert":
                result = self._batch_convert(
                    input_dir, pattern, options, str(base_dir)
                )
            elif operation == "list_files":
                result = self._list_files(input_dir, pattern)
            else:
                return {
                    "success": False,
                    "error": f"不支持的操作: {operation}",
                }
            return {"success": True, "operation": operation, **result}
        except ImportError as e:
            return {
                "success": False,
                "error": f"依赖库缺失: {e}. 请安装对应依赖后重试。",
            }
        except Exception as e:
            return {"success": False, "error": f"批量处理失败: {e}"}

    async def stream(self, inputs, **kwargs):
        yield await self.invoke(inputs, **kwargs)

    @staticmethod
    def _find_files(input_dir: str, pattern: str) -> list[Path]:
        return sorted(Path(input_dir).glob(pattern))

    def _batch_rename(
        self, input_dir: str, pattern: str, options: dict
    ) -> dict:
        prefix = options.get("prefix", "file_")
        start_num = options.get("start_num", 1)
        padding = options.get("padding", 3)
        files = self._find_files(input_dir, pattern)
        renamed = []
        for i, f in enumerate(files):
            new_name = (
                f"{prefix}{str(start_num + i).zfill(padding)}{f.suffix}"
            )
            new_path = f.parent / new_name
            f.rename(new_path)
            renamed.append({"original": f.name, "new": new_name})
        return {"total_files": len(renamed), "renamed": renamed}

    def _merge_excel_files(
        self, input_dir: str, pattern: str, output_dir: str
    ) -> dict:
        from openpyxl import Workbook, load_workbook

        files = self._find_files(input_dir, pattern or "*.xlsx")
        if not files:
            return {"total_files": 0, "error": "未找到匹配文件"}

        merged_wb = Workbook()
        merged_ws = merged_wb.active
        merged_ws.title = "Merged"
        header_written = False
        total_rows = 0

        for f in files:
            wb = load_workbook(str(f), data_only=True)
            ws = wb.active
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue
            if not header_written:
                merged_ws.append(
                    [str(c) if c is not None else "" for c in rows[0]]
                )
                header_written = True
            for row in rows[1:]:
                merged_ws.append(
                    [str(c) if c is not None else "" for c in row]
                )
                total_rows += 1

        output_path = str(Path(output_dir) / "merged_excel.xlsx")
        merged_wb.save(output_path)
        return {
            "total_files": len(files),
            "total_rows": total_rows,
            "path": output_path,
            "exists": os.path.isfile(output_path),
            "size_bytes": (
                os.path.getsize(output_path)
                if os.path.isfile(output_path)
                else 0
            ),
        }

    def _merge_word_files(
        self, input_dir: str, pattern: str, output_dir: str
    ) -> dict:
        from docx import Document

        files = self._find_files(input_dir, pattern or "*.docx")
        if not files:
            return {"total_files": 0, "error": "未找到匹配文件"}

        merged_doc = Document()
        for i, f in enumerate(files):
            if i > 0:
                merged_doc.add_page_break()
            doc = Document(str(f))
            for para in doc.paragraphs:
                merged_doc.add_paragraph(para.text, style=para.style)
            for table in doc.tables:
                rows = len(table.rows)
                cols = len(table.columns)
                new_table = merged_doc.add_table(rows=rows, cols=cols)
                for r in range(rows):
                    for c in range(cols):
                        new_table.rows[r].cells[c].text = (
                            table.rows[r].cells[c].text
                        )

        output_path = str(Path(output_dir) / "merged_word.docx")
        merged_doc.save(output_path)
        return {
            "total_files": len(files),
            "path": output_path,
            "exists": os.path.isfile(output_path),
            "size_bytes": (
                os.path.getsize(output_path)
                if os.path.isfile(output_path)
                else 0
            ),
        }

    def _batch_convert(
        self,
        input_dir: str,
        pattern: str,
        options: dict,
        output_dir: str,
    ) -> dict:
        target_format = options.get("target_format", "pdf")
        files = self._find_files(input_dir, pattern)
        if not files:
            return {"total_files": 0, "error": "未找到匹配文件"}

        converted = []
        failed = []
        for f in files:
            try:
                output_path = self._convert_single(
                    str(f), target_format, output_dir
                )
                if output_path:
                    converted.append(
                        {
                            "source": f.name,
                            "output": Path(output_path).name,
                        }
                    )
                else:
                    failed.append(
                        {"file": f.name, "error": "不支持的转换路径"}
                    )
            except Exception as e:
                failed.append({"file": f.name, "error": str(e)})
        return {
            "total_files": len(files),
            "converted": len(converted),
            "failed": len(failed),
            "converted_files": converted,
            "failed_files": failed,
        }

    @staticmethod
    def _convert_single(
        source_path: str, target_format: str, output_dir: str
    ) -> str | None:
        """Convert a single file, return output path or None."""
        source_ext = Path(source_path).suffix.lower()
        stem = Path(source_path).stem
        ext_map = {
            "word": ".docx",
            "pdf": ".pdf",
            "excel": ".xlsx",
            "csv": ".csv",
            "ppt": ".pptx",
        }
        output_path = str(
            Path(output_dir)
            / f"{stem}_converted{ext_map.get(target_format, f'.{target_format}')}"
        )

        if source_ext in (".xls", ".xlsx") and target_format == "csv":
            import csv

            from openpyxl import load_workbook

            wb = load_workbook(source_path, data_only=True)
            ws = wb.active
            with open(output_path, "w", encoding="utf-8", newline="") as f:
                writer = csv.writer(f)
                for row in ws.iter_rows(values_only=True):
                    writer.writerow(
                        [str(c) if c is not None else "" for c in row]
                    )
            return output_path

        if source_ext == ".csv" and target_format == "excel":
            import csv

            from openpyxl import Workbook

            wb = Workbook()
            ws = wb.active
            with open(source_path, encoding="utf-8") as f:
                reader = csv.reader(f)
                for row in reader:
                    ws.append(row)
            wb.save(output_path)
            return output_path

        if source_ext in (".doc", ".docx") and target_format == "pdf":
            from docx import Document
            from fpdf import FPDF

            doc = Document(source_path)
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            font_name = select_pdf_font(pdf)
            pdf.set_font(font_name, "", 11)
            for para in doc.paragraphs:
                if para.text.strip():
                    pdf.multi_cell(0, 7, para.text)
                    pdf.ln(3)
            pdf.output(output_path)
            return output_path

        return None

    def _list_files(self, input_dir: str, pattern: str) -> dict:
        files = self._find_files(input_dir, pattern or "*")
        file_list = []
        for f in files:
            stat = f.stat()
            file_list.append(
                {
                    "name": f.name,
                    "path": str(f),
                    "size_bytes": stat.st_size,
                    "is_dir": f.is_dir(),
                    "extension": f.suffix.lower(),
                }
            )
        return {"total_files": len(file_list), "files": file_list}
