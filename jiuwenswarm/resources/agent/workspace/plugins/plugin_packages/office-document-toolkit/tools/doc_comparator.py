import os
from pathlib import Path

from openjiuwen.core.foundation.tool import Tool, ToolCard


class DocComparator(Tool):
    """文档比较工具：对比两个Word/Excel/PDF/CSV文档的差异。"""

    def __init__(self) -> None:
        super().__init__(
            ToolCard(
                id="doc_comparator",
                name="doc_comparator",
                description=(
                    "文档比较工具：对比两个Word/Excel/PDF/CSV文档的"
                    "文本差异，输出新增、删除、修改的内容。"
                    "当用户需要比较文档版本差异时调用。"
                ),
                input_params={
                    "type": "object",
                    "properties": {
                        "file_path_1": {
                            "type": "string",
                            "description": "第一个文件路径（旧版）",
                        },
                        "file_path_2": {
                            "type": "string",
                            "description": "第二个文件路径（新版）",
                        },
                        "output_format": {
                            "type": "string",
                            "enum": [
                                "structured",
                                "markdown",
                                "json",
                            ],
                            "description": "输出格式，默认 structured",
                        },
                    },
                    "required": ["file_path_1", "file_path_2"],
                },
            )
        )

    async def invoke(self, inputs, **kwargs):
        file_path_1 = inputs.get("file_path_1", "")
        file_path_2 = inputs.get("file_path_2", "")
        output_format = inputs.get("output_format", "structured")

        if not file_path_1 or not os.path.isfile(file_path_1):
            return {
                "success": False,
                "error": f"文件1不存在: {file_path_1}",
            }
        if not file_path_2 or not os.path.isfile(file_path_2):
            return {
                "success": False,
                "error": f"文件2不存在: {file_path_2}",
            }

        try:
            ext1 = Path(file_path_1).suffix.lower()
            ext2 = Path(file_path_2).suffix.lower()

            if ext1 != ext2:
                return {
                    "success": False,
                    "error": "两个文件类型不一致，无法比较",
                }

            if ext1 in (".doc", ".docx"):
                result = self._compare_word(file_path_1, file_path_2)
            elif ext1 in (".xls", ".xlsx"):
                result = self._compare_excel(
                    file_path_1, file_path_2
                )
            elif ext1 == ".pdf":
                result = self._compare_pdf(file_path_1, file_path_2)
            elif ext1 == ".csv":
                result = self._compare_csv(file_path_1, file_path_2)
            else:
                return {
                    "success": False,
                    "error": f"不支持的文件类型: {ext1}",
                }

            if output_format == "markdown":
                result = self._to_markdown(result)
            elif output_format == "json":
                import json

                result = json.dumps(
                    result, ensure_ascii=False, indent=2
                )

            return {
                "success": True,
                "file_1": file_path_1,
                "file_2": file_path_2,
                "diff": result,
            }
        except ImportError as e:
            return {
                "success": False,
                "error": f"依赖库缺失: {e}. 请安装对应依赖后重试。",
            }
        except Exception as e:
            return {"success": False, "error": f"比较失败: {e}"}

    async def stream(self, inputs, **kwargs):
        yield await self.invoke(inputs, **kwargs)

    @staticmethod
    def _compare_word(path1: str, path2: str) -> dict:
        import difflib

        from docx import Document

        doc1 = Document(path1)
        doc2 = Document(path2)

        paras1 = [p.text for p in doc1.paragraphs if p.text.strip()]
        paras2 = [p.text for p in doc2.paragraphs if p.text.strip()]

        diff = list(
            difflib.unified_diff(paras1, paras2, lineterm="", n=2)
        )

        added = []
        removed = []
        for line in diff:
            if line.startswith("+") and not line.startswith("+++"):
                added.append(line[1:])
            elif line.startswith("-") and not line.startswith("---"):
                removed.append(line[1:])

        return {
            "file_type": "word",
            "paragraphs_1": len(paras1),
            "paragraphs_2": len(paras2),
            "added": added,
            "removed": removed,
            "added_count": len(added),
            "removed_count": len(removed),
        }

    @staticmethod
    def _compare_excel(path1: str, path2: str) -> dict:
        from openpyxl import load_workbook
        from openpyxl.utils import get_column_letter

        wb1 = load_workbook(path1, data_only=True)
        wb2 = load_workbook(path2, data_only=True)

        sheets1 = set(wb1.sheetnames)
        sheets2 = set(wb2.sheetnames)

        added_sheets = sheets2 - sheets1
        removed_sheets = sheets1 - sheets2
        common_sheets = sheets1 & sheets2

        cell_diffs = []
        for sheet_name in common_sheets:
            ws1 = wb1[sheet_name]
            ws2 = wb2[sheet_name]
            max_row = max(ws1.max_row, ws2.max_row)
            max_col = max(ws1.max_column, ws2.max_column)
            for row in range(1, max_row + 1):
                for col in range(1, max_col + 1):
                    val1 = ws1.cell(row=row, column=col).value
                    val2 = ws2.cell(row=row, column=col).value
                    if val1 != val2:
                        cell_ref = (
                            f"{get_column_letter(col)}{row}"
                        )
                        cell_diffs.append(
                            {
                                "sheet": sheet_name,
                                "cell": cell_ref,
                                "old_value": (
                                    str(val1)
                                    if val1 is not None
                                    else ""
                                ),
                                "new_value": (
                                    str(val2)
                                    if val2 is not None
                                    else ""
                                ),
                            }
                        )

        return {
            "file_type": "excel",
            "sheets_1": list(sheets1),
            "sheets_2": list(sheets2),
            "added_sheets": list(added_sheets),
            "removed_sheets": list(removed_sheets),
            "cell_differences": cell_diffs,
            "total_differences": len(cell_diffs),
        }

    @staticmethod
    def _compare_pdf(path1: str, path2: str) -> dict:
        import difflib

        from pypdf import PdfReader

        reader1 = PdfReader(path1)
        reader2 = PdfReader(path2)

        text1 = [
            page.extract_text() or "" for page in reader1.pages
        ]
        text2 = [
            page.extract_text() or "" for page in reader2.pages
        ]

        diff = list(
            difflib.unified_diff(text1, text2, lineterm="", n=2)
        )

        added = []
        removed = []
        for line in diff:
            if line.startswith("+") and not line.startswith("+++"):
                added.append(line[1:])
            elif line.startswith("-") and not line.startswith("---"):
                removed.append(line[1:])

        return {
            "file_type": "pdf",
            "pages_1": len(reader1.pages),
            "pages_2": len(reader2.pages),
            "added": added,
            "removed": removed,
            "added_count": len(added),
            "removed_count": len(removed),
        }

    @staticmethod
    def _compare_csv(path1: str, path2: str) -> dict:
        import csv
        import difflib

        with open(path1, encoding="utf-8") as f:
            rows1 = list(csv.reader(f))
        with open(path2, encoding="utf-8") as f:
            rows2 = list(csv.reader(f))

        diff = list(
            difflib.unified_diff(
                [",".join(r) for r in rows1],
                [",".join(r) for r in rows2],
                lineterm="",
                n=2,
            )
        )

        added = []
        removed = []
        for line in diff:
            if line.startswith("+") and not line.startswith("+++"):
                added.append(line[1:])
            elif line.startswith("-") and not line.startswith("---"):
                removed.append(line[1:])

        return {
            "file_type": "csv",
            "rows_1": len(rows1),
            "rows_2": len(rows2),
            "added": added,
            "removed": removed,
            "added_count": len(added),
            "removed_count": len(removed),
        }

    @staticmethod
    def _to_markdown(result: dict) -> str:
        lines = ["# 文档差异报告", ""]
        lines.append(f"- 文件类型: {result.get('file_type', '')}")
        lines.append("")

        if result.get("added_count") is not None:
            lines.append(
                f"## 新增内容 ({result['added_count']} 项)"
            )
            lines.append("")
            for item in result.get("added", []):
                lines.append(f"+ {item[:100]}")
            lines.append("")
            lines.append(
                f"## 删除内容 ({result['removed_count']} 项)"
            )
            lines.append("")
            for item in result.get("removed", []):
                lines.append(f"- {item[:100]}")
            lines.append("")

        if result.get("cell_differences"):
            lines.append(
                f"## 单元格差异 "
                f"({result['total_differences']} 处)"
            )
            lines.append("")
            for diff in result["cell_differences"]:
                lines.append(
                    f"- **{diff['sheet']}!{diff['cell']}**: "
                    f"{diff['old_value']} -> {diff['new_value']}"
                )
            lines.append("")

        if result.get("added_sheets"):
            lines.append(
                f"## 新增工作表: "
                f"{', '.join(result['added_sheets'])}"
            )
        if result.get("removed_sheets"):
            lines.append(
                f"## 删除工作表: "
                f"{', '.join(result['removed_sheets'])}"
            )

        return "\n".join(lines)
