import os
from pathlib import Path

from openjiuwen.core.foundation.tool import Tool, ToolCard


class MailMergeTool(Tool):
    """邮件合并工具：用Excel/CSV数据批量填充Word模板，生成多份文档。"""

    def __init__(self) -> None:
        super().__init__(
            ToolCard(
                id="mail_merge_tool",
                name="mail_merge_tool",
                description=(
                    "邮件合并工具：用Excel/CSV数据批量填充Word模板中的"
                    "占位符（如{{name}}），生成多份个性化文档。"
                    "当用户需要批量生成合同/邀请函/工资条等时调用。"
                ),
                input_params={
                    "type": "object",
                    "properties": {
                        "template_path": {
                            "type": "string",
                            "description": "Word模板文件路径（.docx），含占位符如{{name}}",
                        },
                        "data_path": {
                            "type": "string",
                            "description": "数据源文件路径（.xlsx或.csv）",
                        },
                        "output_subdir": {
                            "type": "string",
                            "description": "输出子目录名，默认为 mail_merge_output",
                        },
                        "filename_pattern": {
                            "type": "string",
                            "description": "输出文件名模式，如{{name}}_合同，默认用序号",
                        },
                    },
                    "required": ["template_path", "data_path"],
                },
            )
        )

    async def invoke(self, inputs, **kwargs):
        template_path = inputs.get("template_path", "")
        data_path = inputs.get("data_path", "")
        output_subdir = inputs.get(
            "output_subdir", "mail_merge_output"
        )
        filename_pattern = inputs.get("filename_pattern", "")

        if not template_path or not os.path.isfile(template_path):
            return {
                "success": False,
                "error": f"模板文件不存在: {template_path}",
            }
        if not data_path or not os.path.isfile(data_path):
            return {
                "success": False,
                "error": f"数据源文件不存在: {data_path}",
            }

        from openjiuwen.core.sys_operation.cwd import get_cwd

        base_dir = Path(get_cwd()) / output_subdir
        base_dir.mkdir(parents=True, exist_ok=True)

        try:
            data = self._load_data(data_path)
            if not data:
                return {"success": False, "error": "数据源为空"}

            generated = []
            failed = []
            for i, row in enumerate(data):
                try:
                    doc = self._fill_template(template_path, row)
                    if filename_pattern:
                        filename = self._apply_pattern(
                            filename_pattern, row, i
                        )
                    else:
                        filename = f"document_{i + 1}"
                    output_path = base_dir / f"{filename}.docx"
                    doc.save(str(output_path))
                    generated.append(
                        {
                            "index": i + 1,
                            "filename": output_path.name,
                            "path": str(output_path),
                            "size_bytes": output_path.stat().st_size,
                        }
                    )
                except Exception as e:
                    failed.append(
                        {"index": i + 1, "error": str(e)}
                    )

            return {
                "success": True,
                "total_records": len(data),
                "generated": len(generated),
                "failed": len(failed),
                "generated_files": generated[:20],
                "failed_records": failed[:20],
                "output_dir": str(base_dir),
            }
        except ImportError as e:
            return {
                "success": False,
                "error": f"依赖库缺失: {e}. 请安装 python-docx 和 openpyxl 后重试。",
            }
        except Exception as e:
            return {"success": False, "error": f"邮件合并失败: {e}"}

    async def stream(self, inputs, **kwargs):
        yield await self.invoke(inputs, **kwargs)

    @staticmethod
    def _load_data(file_path: str) -> list[dict]:
        ext = Path(file_path).suffix.lower()
        if ext == ".csv":
            import csv

            with open(file_path, encoding="utf-8") as f:
                reader = csv.DictReader(f)
                return list(reader)
        from openpyxl import load_workbook

        wb = load_workbook(file_path, data_only=True)
        ws = wb.active
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            return []
        headers = [
            str(h) if h else f"col_{i}"
            for i, h in enumerate(rows[0])
        ]
        return [dict(zip(headers, row)) for row in rows[1:]]

    @staticmethod
    def _fill_template(template_path: str, data: dict):
        import re

        from docx import Document

        doc = Document(template_path)

        def replace_placeholders(text: str) -> str:
            def replacer(m):
                key = m.group(1).strip()
                val = data.get(key, "")
                return str(val) if val is not None else ""

            return re.sub(r"\{\{(.+?)\}\}", replacer, text)

        for para in doc.paragraphs:
            for run in para.runs:
                run.text = replace_placeholders(run.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = replace_placeholders(run.text)

        for section in doc.sections:
            for para in section.header.paragraphs:
                for run in para.runs:
                    run.text = replace_placeholders(run.text)
            for para in section.footer.paragraphs:
                for run in para.runs:
                    run.text = replace_placeholders(run.text)

        return doc

    @staticmethod
    def _apply_pattern(
        pattern: str, data: dict, index: int
    ) -> str:
        import re

        def replacer(m):
            key = m.group(1).strip()
            if key == "index":
                return str(index + 1)
            val = data.get(key, "")
            return str(val) if val is not None else ""

        result = re.sub(r"\{\{(.+?)\}\}", replacer, pattern)
        result = re.sub(r'[<>:"/\\|?*]', "_", result)
        return result
