import os
import re
from pathlib import Path

from openjiuwen.core.foundation.tool import Tool, ToolCard


class Redactor(Tool):
    """文档脱敏工具：自动识别并遮盖文档中的敏感信息。"""

    def __init__(self) -> None:
        super().__init__(
            ToolCard(
                id="redactor",
                name="redactor",
                description=(
                    "文档脱敏工具：自动识别并遮盖Word/Excel/PDF/TXT中的"
                    "身份证号、手机号、邮箱、银行卡号等敏感信息。"
                    "当用户需要脱敏文档时调用。"
                ),
                input_params={
                    "type": "object",
                    "properties": {
                        "file_path": {
                            "type": "string",
                            "description": "要脱敏的文件路径",
                        },
                        "patterns": {
                            "type": "array",
                            "items": {"type": "string"},
                            "description": "自定义脱敏正则模式，默认全部启用",
                        },
                        "output_subdir": {
                            "type": "string",
                            "description": "输出子目录名，默认为 redacted",
                        },
                    },
                    "required": ["file_path"],
                },
            )
        )

    async def invoke(self, inputs, **kwargs):
        file_path = inputs.get("file_path", "")
        custom_patterns = inputs.get("patterns", [])
        output_subdir = inputs.get("output_subdir", "redacted")

        if not file_path or not os.path.isfile(file_path):
            return {"success": False, "error": f"文件不存在: {file_path}"}

        from openjiuwen.core.sys_operation.cwd import get_cwd

        base_dir = Path(get_cwd()) / output_subdir
        base_dir.mkdir(parents=True, exist_ok=True)

        try:
            ext = Path(file_path).suffix.lower()
            patterns = self._get_patterns(custom_patterns)

            if ext in (".doc", ".docx"):
                output_path, stats = self._redact_word(
                    file_path, patterns, str(base_dir)
                )
            elif ext in (".xls", ".xlsx"):
                output_path, stats = self._redact_excel(
                    file_path, patterns, str(base_dir)
                )
            elif ext == ".pdf":
                output_path, stats = self._redact_pdf(
                    file_path, patterns, str(base_dir)
                )
            elif ext in (".txt", ".csv", ".md"):
                output_path, stats = self._redact_text_file(
                    file_path, patterns, str(base_dir)
                )
            else:
                return {
                    "success": False,
                    "error": f"不支持的文件类型: {ext}",
                }

            return {
                "success": True,
                "file_path": file_path,
                "path": output_path,
                "exists": os.path.isfile(output_path),
                "size_bytes": (
                    os.path.getsize(output_path)
                    if os.path.isfile(output_path)
                    else 0
                ),
                "redactions": stats,
            }
        except ImportError as e:
            return {
                "success": False,
                "error": f"依赖库缺失: {e}. 请安装对应依赖后重试。",
            }
        except Exception as e:
            return {"success": False, "error": f"脱敏失败: {e}"}

    async def stream(self, inputs, **kwargs):
        yield await self.invoke(inputs, **kwargs)

    @staticmethod
    def _get_patterns(
        custom_patterns: list,
    ) -> list[tuple[re.Pattern, str]]:
        """Return list of (compiled_regex, replacement)."""
        default = [
            (
                r"\b1[3-9]\d{9}\b",
                "1**********",
            ),
            (
                r"\b\d{6}(?:19|20)\d{2}(?:0[1-9]|1[0-2])"
                r"(?:0[1-9]|[12]\d|3[01])\d{3}[\dXx]\b",
                "******************",
            ),
            (
                r"\b[\w.+-]+@[\w-]+\.[\w.-]+\b",
                "***@***.***",
            ),
            (
                r"\b\d{4}[-\s]?\d{4}[-\s]?\d{4}[-\s]?\d{4}\b",
                "****-****-****-****",
            ),
            (
                r"\b\d{3}[-]?\d{8}\b",
                "***-********",
            ),
        ]
        result = [
            (re.compile(p), r) for p, r in default
        ]
        for p in custom_patterns:
            result.append((re.compile(p), "***"))
        return result

    @staticmethod
    def _redact_text(
        text: str, patterns: list
    ) -> tuple[str, dict]:
        stats: dict[str, int] = {}
        for regex, replacement in patterns:
            count = len(regex.findall(text))
            if count > 0:
                stats[regex.pattern[:40]] = count
                text = regex.sub(replacement, text)
        return text, stats

    def _redact_word(
        self, file_path: str, patterns: list, output_dir: str
    ) -> tuple:
        from docx import Document

        doc = Document(file_path)
        all_stats: dict[str, int] = {}

        for para in doc.paragraphs:
            for run in para.runs:
                text, stats = self._redact_text(
                    run.text, patterns
                )
                run.text = text
                all_stats.update(stats)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            text, stats = self._redact_text(
                                run.text, patterns
                            )
                            run.text = text
                            all_stats.update(stats)

        output_path = str(
            Path(output_dir)
            / f"{Path(file_path).stem}_redacted.docx"
        )
        doc.save(output_path)
        return output_path, all_stats

    def _redact_excel(
        self, file_path: str, patterns: list, output_dir: str
    ) -> tuple:
        from openpyxl import load_workbook

        wb = load_workbook(file_path)
        all_stats: dict[str, int] = {}

        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value and isinstance(
                        cell.value, str
                    ):
                        text, stats = self._redact_text(
                            cell.value, patterns
                        )
                        cell.value = text
                        all_stats.update(stats)

        output_path = str(
            Path(output_dir)
            / f"{Path(file_path).stem}_redacted.xlsx"
        )
        wb.save(output_path)
        return output_path, all_stats

    def _redact_pdf(
        self, file_path: str, patterns: list, output_dir: str
    ) -> tuple:
        from pypdf import PdfReader, PdfWriter

        reader = PdfReader(file_path)
        writer = PdfWriter()
        all_stats: dict[str, int] = {}

        for page in reader.pages:
            text = page.extract_text() or ""
            _, stats = self._redact_text(text, patterns)
            all_stats.update(stats)
            writer.add_page(page)

        output_path = str(
            Path(output_dir)
            / f"{Path(file_path).stem}_redacted.pdf"
        )
        with open(output_path, "wb") as f:
            writer.write(f)
        return output_path, all_stats

    def _redact_text_file(
        self, file_path: str, patterns: list, output_dir: str
    ) -> tuple:
        ext = Path(file_path).suffix.lower()
        with open(file_path, encoding="utf-8") as f:
            content = f.read()

        redacted, stats = self._redact_text(content, patterns)

        output_path = str(
            Path(output_dir)
            / f"{Path(file_path).stem}_redacted{ext}"
        )
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(redacted)
        return output_path, stats
